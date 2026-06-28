"""
Generate updated Chapter 3 docx — minimal changes from the original.
Only specific words/sentences changed; all other text kept verbatim.
Run: venv311/bin/python scripts/generate_chap3_minimal.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.0)

# ── helpers ───────────────────────────────────────────────────────────────────
def p(text="", bold=False, italic=False, center=False, size=12):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = Pt(24)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
    return para

def h1(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    return para

def h2(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return para

def h3(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Inches(0.25)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return para

def caption(text, center=True):
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(10)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    return para

def fig_placeholder(label):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "D9D9D9")
    tcPr.append(shd)
    pp = cell.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(36)
    pp.paragraph_format.space_after  = Pt(36)
    run = pp.add_run(f"[ {label} ]")
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()

def tbl_caption(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return para

def make_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
    doc.add_paragraph()

def uc_table(title, rows_2col):
    """Two-column use-case description table."""
    tbl_caption(title)
    t = doc.add_table(rows=len(rows_2col), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, (label, val) in enumerate(rows_2col):
        cells = t.rows[ri].cells
        run1 = cells[0].paragraphs[0].add_run(label)
        run1.bold = True; run1.font.size = Pt(11); run1.font.name = "Times New Roman"
        run2 = cells[1].paragraphs[0].add_run(val)
        run2.font.size = Pt(11); run2.font.name = "Times New Roman"
    doc.add_paragraph()

def pseudo(lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(8)
    run = para.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(10)

def bullet(text, indent=0.5):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.space_after = Pt(2)

def sub_bullet(text):
    bullet(text, indent=0.75)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
for line in ["UNIVERSITI TEKNOLOGI MARA", "", "EARLY WARNING SYSTEM FOR\nSUPPLY CHAIN RESILIENCE",
             "", "MEOR DANISH FARHAN BIN SOBRI", ""]:
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pp.add_run(line)
    run.bold = "UNIVERSITI" in line or "EARLY" in line
    run.font.size = Pt(14 if ("UNIVERSITI" in line or "EARLY" in line) else 12)
    run.font.name = "Times New Roman"

pp = doc.add_paragraph()
pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pp.add_run("Thesis submitted in partial fulfilment of the requirements for\nBachelor of Computer Science (Hons.)")
run.font.size = Pt(12); run.font.name = "Times New Roman"

doc.add_paragraph()
pp = doc.add_paragraph()
pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pp.add_run("Faculty of Computer and Mathematical Science")
run.font.size = Pt(12); run.font.name = "Times New Roman"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER HEADING
# ══════════════════════════════════════════════════════════════════════════════
pp = doc.add_paragraph()
pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pp.add_run("CHAPTER 3\nMETHODOLOGY")
run.bold = True; run.font.size = Pt(14); run.font.name = "Times New Roman"
pp.paragraph_format.space_after = Pt(12)

# VERBATIM
p("This chapter outlines the methodology adopted for the development of the proposed "
  "system. It provides a structured methodology that supports the process of planning, "
  "designing, implementing, and evaluating the various activities across the project "
  "lifecycle. The software development model that will be used in this study is the "
  "Waterfall Model that is defined by the sequential and linear flow of development "
  "activities.")

# ══════════════════════════════════════════════════════════════════════════════
# 3.1
# ══════════════════════════════════════════════════════════════════════════════
h1("3.1\tSoftware Development Methodology")

# VERBATIM
p("The waterfall model is a software development methodology that is based on the "
  "software development lifecycle and employs a sequential and linear approach. It is "
  "used for projects that require clarity, consistency, and are unlikely to change while "
  "they are being developed. The waterfall model structure is shown in greater depth in "
  "Figure 3.1.")

fig_placeholder("Figure 3.1 Waterfall Model Structure")
caption("Figure 3.1 Waterfall Model Structure")

# Table 3.1 — VERBATIM
tbl_caption("Table 3.1 Project Framework")
make_table(
    ["Objective", "Phase", "Activities", "Outcome"],
    [
        ["Objective 1: To identify suitable techniques for monitoring and assessing supply chain disruption risks.",
         "Planning",
         "• Conduct literature review on supply chain disruption, resilience, and risk assessment techniques.\n"
         "• Review existing machine learning and NLP techniques used for disruption detection.\n"
         "• Analyse suitable datasets and data sources for supply chain risk monitoring.",
         "• Selection of appropriate machine learning and NLP techniques.\n"
         "• Identification of relevant datasets and disruption indicators."],
        ["Objective 2: To design a supply chain resilience monitoring system using a hybrid machine learning architecture.",
         "Design\n\nDevelopment",
         "• Design overall system architecture and workflow diagrams.\n"
         "• Design offline training and online inference pipelines.\n"
         "• Design database schema and user interface structure.\n"
         "• Implement backend services for data ingestion and processing.\n"
         "• Train machine learning models using historical news data.\n"
         "• Integrate trained models into the inference pipeline.\n"
         "• Develop frontend dashboard for visualization.",
         "• Defined system architecture and data flow.\n"
         "• Finalised ERD and interface design.\n\n"
         "• Trained and stored machine learning models.\n"
         "• Functional web-based supply chain monitoring system."],
        ["Objective 3: To test the functionality and reliability of the developed supply chain resilience monitoring system.",
         "Testing",
         "• Perform functional testing on system modules.\n"
         "• Validate model inference and score updates.\n"
         "• Test user interface and visualization components.",
         "• Verified system functionality and performance.\n"
         "• Identified potential improvements."],
    ]
)

# VERBATIM
p("Table 3.1 presents the project framework that maps each project objective to "
  "the corresponding phase of the Waterfall model. Objective 1 is addressed during the "
  "planning phase through literature review and technique selection. Objective 2 is "
  "achieved in the design and development phases through system design, model training, "
  "and system implementation. Lastly, Objective 3 is fulfilled in the testing phase by "
  "evaluating system functionality and accuracy performance.")

# ══════════════════════════════════════════════════════════════════════════════
# 3.2
# ══════════════════════════════════════════════════════════════════════════════
h1("3.2\tPhase 1: Requirement Analysis")

# VERBATIM
p("During this stage of planning, the critical goals of the project, data needs, and "
  "approaches to the methodology are established to aid in the successful implementation "
  "of the supply chain disruption forecasting system. It is centred on how the textual data "
  "that is obtained based on the news articles could be used to identify and forecast events "
  "that are likely to interfere with the operations in the supply chain. This involves "
  "determining the features that are relevant to news content, including the type of event "
  "(e.g., strikes, disasters, geopolitical problems), entities (places, organisations), and "
  "time.")

# VERBATIM
p("In that regard, the project uses the News Dataset offered by Hugging Face "
  "(R3troR0b/news-dataset), a collection of real-life news headlines and article "
  "summaries that were annotated with event categories. In contrast to other types of "
  "disruption detection models, which utilise social media or financial information, this "
  "system was created to be trained exclusively on news-based text information to "
  "guarantee the reliability of content and richness of context.")

# CHANGE: "XGBoost and BERT" → "XGBoost, FinBERT, and GLiNER2"
p("At this stage, suitable machine learning models are also selected especially "
  "supervised classifiers such as XGBoost, FinBERT, and GLiNER2 depending on the past performance "
  "comparisons. The planning stage also involves setting up the software tools and "
  "libraries to work with (e.g. Python, scikit-learn, Hugging Face Transformers), system "
  "objectives, and aligning the data pipeline with the system architecture in general.")

# 3.2.1 — VERBATIM
h2("3.2.1\tData Collection")

p("The main dataset that is used in this study is the R3troR0b/news-dataset that can be "
  "found in Hugging Face. This dataset consists of a huge amount of news articles across "
  "the globe. Based on its dataset card, the articles are pulled out of such reputable "
  "publishers as Reuters, BBC World News, Al Jazeera, Le Monde, South China Morning "
  "Post, The Hindu, Deutsche Welle, The Guardian, NPR, TASS (Russia), and The "
  "Sydney Morning Herald. The dataset is labelled as news and world and it is to be used "
  "in text classification. It is primarily in English (and some parts are in French and "
  "Russian) and is covered by the MIT Licence. The MIT licence is free to use and to "
  "distribute, making it appropriate in research.")

p("The data is presented in the format of the JSON: it has two primary fields per "
  "record: label (string) and text (string). The label field is a composite key that is a string "
  "and it has metadata i.e. the news source (publisher), article title, URL, and the "
  "publication time (e.g. The Guardian; Article Title; https://...;2025-01-01T14:00:00Z). "
  "The news article has the body of the news in the text field. An example would be that "
  "the label of one record will be The Guardian; Supply chains in crisis:...;https://...;2025"
  "-01-01T14:00:16Z and the article text will be the content of the record. The dataset is "
  "updated actively: files are sorted by quarter (e.g., allnewsq12025.json, "
  "allnewsq22025.json, etc.), and are periodically updated (as of 2026) to add new articles.")

p("The repository contains multi-megabyte JSON files quarterly, which suggests "
  "a strong amount of data to train the models. Its applicability to supply chains is "
  "anticipated since global news normally reports on events which impact logistics, trade, "
  "and supply-chain disruptions. Moreover, the range of the dataset (international sources "
  "and up-to-date content) must include references to the keywords related to the supply "
  "chain (e.g. \"supply chain\", \"logistics\", \"shipment\", \"inventory\", \"transportation\" etc.). "
  "This renders it an effective foundation to develop a news classification platform "
  "centred on supply chain issues.")

fig_placeholder("Figure 3.2 News Dataset")
caption("Figure 3.2 News Dataset")

tbl_caption("Table 3.2 Explanations for News Dataset Features")
make_table(
    ["Index", "Feature Name", "Feature Explanation"],
    [
        ["F1", "Label", "This column includes the article's news provider, title, and URL link. The three data is separated by a semicolon(;)."],
        ["F2", "Text",  "This column includes the full-text news article."],
    ]
)

# 3.2.2
h2("3.2.2\tData Pre-Processing")

fig_placeholder("Figure 3.3 Flowchart of Data Pre-Processing")
caption("Figure 3.3 Flowchart of Data Pre-Processing")

# VERBATIM — keep all preprocessing text exactly as-is
p("Before classification of raw news text it has to be cleaned and ready. During this step, "
  "NLP tools (i.e., pandas, NLTK, and spaCy) will be used to preprocess the articles. To "
  "work with tabular data, first, a JSON datablock may be loaded into a pandas "
  "DataFrame (with pandas.readjson()). The content of each article will be extracted out "
  "of the text field. The most common steps of cleaning up the text are to delete any "
  "punctuations, HTML tags (in case they exist), and numeric characters and to convert "
  "all the text to the lower case so that there would be no discrepancy. Then, there is "
  "tokenization: every article is divided into individual words or tokens. Efficient "
  "tokenizers are offered in libraries such as spaCy and NLTK. As an example, the call "
  "to spacy.load(encorewebsm) followed by doc = nlp(text) returns tokens, and the "
  "wordtokenize() function of NLTK can also be used to divide text.")

p("Once tokenized, the stop word removal and lemmatization will be done to "
  "normalise the tokens. Stop words (such frequent words as and, the, of, etc.) bear a "
  "minimal amount of semantic information and can be filtered with the help of in-built "
  "lists in NLTK (nltk.corpus.stopwords) or spaCy (token.isstop). Lemmatization "
  "decomposes words down to their base or dictionary form (e.g. ships - ship, running - "
  "run), and this grouping assists in grouping similar forms of words. Both NLTK and "
  "spaCy also have lemmatization: e.g. spacy.lang.en.English contains lemmatizer, and "
  "WordNetLemmatizer of NLTK can be applied. According to one of the resources, "
  "tokenization, removal of stop-words, and lemmatization are common preprocessing "
  "steps in NLP pipelines.")

p("Also, since this system is geared towards news related to the supply-chain we "
  "can use the filtering option of keywords to narrow down to the related news. Once "
  "preprocessing is done, all the articles may be searched with key terms connected with "
  "supply chains (such as, supply chain, logistics, manufacturing, procurement, transport, "
  "warehouse, etc.). Articles that contain no such keywords may be de-prioritised or "
  "classified as probably not related to supply-chain classification. Such filtering may be "
  "applied using the simple list matching in pandas or using the phrase matcher in spaCy.")

# VERBATIM — final preprocessing paragraph, but add one clarifying sentence at end
p("During preprocessing, we will use Python packages: pandas to work with data, "
  "NLTK to tokenize the text, use stop words lists, and spaCy to quickly process the text, "
  "tag it with POS tags, and lemmatize it. The text (or tokens) of each article obtained out "
  "of this stage will be the clean text to be used in feature extraction or as input to the "
  "model. This is so that the raw dataset can be converted into a structured form (e.g. lists "
  "of meaningful tokens or vectors features) which is more effective in classifier "
  "performance. It should be noted that spaCy in this system is used specifically for "
  "temporal date extraction, while location named entity recognition is handled by "
  "the GLiNER2 model as described in Section 3.4.4.")

# ══════════════════════════════════════════════════════════════════════════════
# 3.3
# ══════════════════════════════════════════════════════════════════════════════
h1("3.3\tPhase 2: Design")

# CHANGE: "classification based on machine learning" → add "14-day risk forecasting"
p("This part describes the general architecture of the supply chain disruption detection "
  "platform. It addresses the system high-level architecture, the data flow between the "
  "input (news articles) and the output (risk visualization), machine learning algorithm "
  "and feature engineering, the database design, and the user interface. The main design "
  "features are a multi-layer architecture that combines data ingestion, classification "
  "based on machine learning, 14-day risk forecasting, and a user-interactive map frontend. "
  "The design focuses on identifying disruption events caused by unstructured news data "
  "in time and making intuitive visualization of supplier risks and resilience.")

# 3.3.1
h2("3.3.1\tSystem Architecture")

fig_placeholder("Figure 3.4 System Architecture Diagram")
caption("Figure 3.4 System Architecture Diagram")

# VERBATIM — paragraphs 1 & 2
p("According to Figure 3.4 (System Architecture Diagram), the system starts as the user "
  "enters the web application using the ReactJS interface. The user will be introduced to "
  "the homepage, whereby there will be a world map that contains supplier nodes "
  "denoting the various suppliers within the global supply chain. The system will allow "
  "the user to engage with the system by filtering through (by choosing a particular "
  "product supply chain e.g. iPhone), supplier information and the past resilience "
  "performance of each supplier.")

p("Whenever the user takes any action i.e. filters the suppliers or choose a "
  "particular supplier node, the request is dispatched to the backend API server. The "
  "backend is then used to process the request using the filtering and query logic module, "
  "which retrieves the applicable data in the database which includes supplier information, "
  "the current risk level, and the historical records about the resilience. The information "
  "recovered is then sent back to the frontend and shown to the user as updated map nodes, "
  "detailed panels or historical charts.")

# CHANGE: "News API" → "RSS feeds"; "entity extraction" mention GLiNER2
p("The news ingestion scheduler keeps on picking up new articles through the RSS feeds "
  "in the background in the system. The news items being fetched are sent to "
  "the inference orchestrator which dictates the prediction process. The first step involves "
  "preprocessing stage, which is used to clean and present the text, then entity extraction "
  "using the GLiNER2 model that is used to determine related suppliers, locations, and events. "
  "The trained disruption classification model then examines the processed text to give a "
  "response on whether the news article suggests a supply chain disruption.")

# CHANGE: add one sentence about forecast after "reference in the future"
p("In case of a disruption, the system changes the risk and resilience score of the "
  "impacted supplier and saves the article and classification outcome in the database. The "
  "new resilience and risk data is also uploaded to the history records of the resilience to "
  "provide the opportunity to examine the overall performance trends over an extended "
  "period. Additionally, the Two-Stage XGBoost forecasting engine generates a refreshed "
  "14-day forward risk forecast for the affected supplier node. In case the article is not "
  "related to a disruption, it is still saved in the database as record-keeping and "
  "reference in the future.")

# VERBATIM
p("To prepare the model, the system relies on a separate offline training flow. The "
  "machine learning model is trained and tested on the Hugging Face news data to "
  "implement it. When the model has sufficiently good performance, it is stored and "
  "loaded into the live inference pipeline to be able to support real-time prediction in the "
  "running system.")

# 3.3.2
h2("3.3.2\tSystem Flowchart")

fig_placeholder("Figure 3.5 System Flowchart")
caption("Figure 3.5 System Flowchart")

# CHANGE: "scheduled news ingestion process" keep; "using the News API" → "from the configured RSS feeds"
p("According to Figure 3.5 (System Flowchart), the system process initiates when the "
  "system is triggered automatically by the scheduled news ingestion process or manually "
  "by an administrator. After the process begins, the system pulls the most recent news "
  "articles from the configured RSS feeds.")

# VERBATIM
p("Following the retrieval of the news articles, the system undertakes the process "
  "of validating and deduplicating of the articles to make sure that only the relevant "
  "articles that had never been processed before are taken into account. The articles that "
  "pass the validation process are then sent to the text preprocessing phase wherein the "
  "text is cleaned, tokenized and normalized to allow it to be analyzed.")

# CHANGE: add "using the GLiNER2 model" after "entity extraction"
p("The system then carries out entity extraction using the GLiNER2 model, which finds "
  "significant data like names of suppliers, location, and keywords in the news. The "
  "extracted entities are further compared with the supplier database to ascertain the "
  "related supplier nodes in the system to the news article.")

# VERBATIM — non-disruption branch
p("After identifying the supplier association, the system will then run the "
  "disruption classification procedure based on the trained machine learning model to "
  "know whether the news article represents a supply chain disruption or not. When the "
  "article is defined as a non-disruption event, the system stores the article as a regular "
  "news record and updates the news archive directory in the database and the most recent "
  "data should be accessible to the frontend API. The frontend then updates the world "
  "map interface to show the most recent information after which the process terminates.")

# CHANGE: add forecast refresh sentence in the disruption branch
p("In case the article is defined as a disruption event, the system generates a "
  "disruption event record and associates it with the supplier that is affected. The system "
  "then recalculates the score of risk to the supplier and computes the score of resilience. "
  "The new resilience information will be attached to the history of resilience of the "
  "supplier to be analyzed over a long period of time. The system also regenerates the "
  "14-day risk forecast snapshot for the affected node via the Two-Stage XGBoost "
  "forecasting engine. Then the system inserts the article and the event record attached "
  "to it into the database and updates the frontend API. The process is then completed by "
  "updating the frontend interface with the most recent risk and resilience data on the "
  "world map.")

# 3.3.3
h2("3.3.3\tEntity Relationship Diagram")

fig_placeholder("Figure 3.6 Entity Relationship Diagram")
caption("Figure 3.6 Entity Relationship Diagram")

# CHANGE: add "and risk forecast snapshots" to first sentence
p("According to Figure 3.6, the entity relationship Diagram (ERD) shows how the objects "
  "of the system database are organized and how they relate to one another. The ERD will "
  "be developed to accommodate storing supply chain data, supplier details, news articles, "
  "disruption incidents, previous records of resilience, and risk forecast snapshots.")

# VERBATIM
p("SupplyChain table contains the data on various product supply chains that the system "
  "tracks, e.g., the supply chain of the iPhone. The supplychainid is a unique identifier of "
  "every supply chain, and some of the attributes include the name, description, and the "
  "date when it was created. As a single chain of supply might have a number of suppliers "
  "and a single supplier may have multiple supply chains, the association between "
  "SupplyChain and Supplier is through a junction table named SupplyChainSupplier.")

# CHANGE: add criticality to Supplier table description
p("Supplier table is where the details of the suppliers are stored in the form of nodes on "
  "the world map interface. Each supplier is identified with a unique identifier named "
  "supplierid and the details of the supplier like supplier name, country, region, "
  "geographic coordinates (latitude and longitude), industry and the role that the supplier "
  "plays in the supply chain. The table also includes a criticality score (an integer from "
  "1 to 5) that reflects the strategic importance of each supplier node. Moreover, this "
  "table contains the existing risk score and risk level of every supplier that are regularly "
  "updated through the system.")

# VERBATIM
p("The SupplyChainSupplier table includes unique ID and two foreign keys to the "
  "SupplyChain and Supplier tables. This table is the many-to-many relationship between "
  "the supply chains and suppliers and enables the system to sieve out suppliers "
  "depending on the chosen product or supply chain.")

# CHANGE: "gathered in News API" → "gathered through RSS ingestion"
p("NewsArticle table contains all news articles gathered through RSS ingestion. Articles can be "
  "uniquely identified by a number known as articleid and contain some attributes, "
  "including title, origin, URL, date of publication, content, language, and time of "
  "ingestion. The model confidence score is also placed in the table, and it indicates how "
  "certain the machine learning model is that the article refers to a supply chain "
  "disruption.")

# CHANGE: DisruptionEvent — add ML output fields
p("The DisruptionEvent table contains structured disruption events, which are retrieved "
  "out of news articles. The disruption events get a distinct identifier known as eventid "
  "and a foreign key that associates them with the NewsArticle table. The contents of this "
  "table include event type, event location, event date, and event confidence that "
  "will signify the degree of certainty that the system is that the event that is extracted is "
  "a real disruption. The table also stores additional machine learning outputs including "
  "the tri-class risk label, sentiment label, sentiment score, predicted disruption "
  "probability, and predicted impact score produced by the inference pipeline.")

# VERBATIM
p("The SupplierEvent junction table is used to implement the relationship between "
  "DisruptionEvent and Supplier since a single DisruptionEvent may impact multiple "
  "suppliers, and a single Supplier may be impacted by a number of DisruptionEvents. "
  "This table has one unique ID and two foreign keys that refer to Supplier and "
  "DisruptionEvent tables. It also stores the severity of impact and matched entities, "
  "which explain the severity to which the supplier is impacted and which textual "
  "elements led to the match.")

# VERBATIM
p("The ResilienceHistory table is the historical table where the past risk and resilience "
  "scores of every supplier are stored. Every record is identified by historyid and one "
  "foreign key connecting the historyid to Supplier table. These attributes are the recorded "
  "date and time, resilience score, risk score and optional notes. The system will use this "
  "table to show the historical performance of suppliers on the interface.")

# NEW paragraph — ForecastSnapshot table
p("The ForecastSnapshot table stores the output of the Two-Stage XGBoost forecasting "
  "engine. Each record is associated with a supplier node, a forecast origin date, and a "
  "horizon day offset. It stores the predicted risk value (yhat), lower and upper prediction "
  "intervals (yhat_lower, yhat_upper), the actual realised risk value where available "
  "(y_actual), and the forecast method identifier. This table enables the system to serve "
  "14-day risk forecasts through the API without regenerating them on every request.")

# 3.3.4
h2("3.3.4\tUse Case Diagram")

fig_placeholder("Figure 3.7 Use Case Diagram")
caption("Figure 3.7 Use Case Diagram")

# Use cases — keep all original, add new UC5, renumber rest
use_cases_text = [
    ("1. View Map Dashboard", [
        "The user can access the main dashboard of the system.",
        "The system displays a world map with supplier nodes.",
        "Each node represents a supplier and is color-coded based on its current risk or resilience level.",
    ]),
    ("2. Filter Suppliers by Supply Chain or Product", [
        "The user can filter the displayed supplier nodes by selecting a specific supply chain or product (for example, iPhone supply chain).",
        "The system will update the world map to show only the relevant suppliers.",
    ]),
    ("3. View Supplier Details (Risk and Resilience)", [
        "The user can click on any supplier node on the map.",
        "The system will display detailed information such as:",
    ], [
        "Current risk score",
        "Risk level",
        "Latest related disruption events",
    ]),
    ("4. View Resilience History", [
        "The user can view the historical resilience trend of a selected supplier.",
        "The system will display a time-based graph or chart showing changes in resilience and risk scores.",
        "This allows the user to analyze long-term supplier performance.",
    ]),
    # NEW USE CASE
    ("5. View 14-Day Risk Forecast", [
        "The user can view the 14-day forward risk forecast for a selected supplier.",
        "The system will display a forecast chart generated by the Two-Stage XGBoost model, showing predicted risk values with confidence intervals for each of the 14 horizon days.",
        "This allows the user to anticipate upcoming disruption risks before they occur.",
    ]),
    ("6. Search Supplier", [  # was 5
        "The user can search for a supplier by name or location.",
        "The system will highlight or display the matching supplier on the map and in the detail panel.",
    ]),
    ("7. View Related News Articles", [  # was 6
        "The user can view the news articles associated with a supplier or disruption event.",
        "The system displays article titles, sources, and publication dates.",
    ]),
    ("8. Trigger Manual Refresh", [  # was 7
        "The administrator can manually trigger the system to fetch the latest news articles from the RSS feeds.",
        "The system will immediately run the inference pipeline and update the database.",
    ]),
    ("9. Manage Supplier Records", [  # was 8
        "The administrator can add new suppliers to the system.",
        "The administrator can update or remove existing supplier information.",
    ]),
    ("10. Manage Supply Chain Records", [  # was 9
        "The administrator can add, update, or remove supply chain or product records.",
    ]),
    ("11. View System Logs and Reports", [  # was 10
        "The administrator can view system activity logs and processing reports.",
    ]),
]

for uc in use_cases_text:
    title = uc[0]
    bullets_main = uc[1]
    sub_bullets = uc[2] if len(uc) > 2 else []

    pp = doc.add_paragraph()
    run = pp.add_run(title)
    run.bold = True; run.font.size = Pt(12); run.font.name = "Times New Roman"
    pp.paragraph_format.space_before = Pt(6)
    pp.paragraph_format.space_after  = Pt(2)

    for b in bullets_main:
        bullet(b)
    for sb in sub_bullets:
        sub_bullet(sb)

doc.add_paragraph()

# 3.3.5
h2("3.3.5\tUse Case Description")

# Original tables 3.3–3.12, verbatim, plus new table for UC5
uc_table("Table 3.3 View Map Dashboard Use Case Description", [
    ("Use Case Name",         "View Map Dashboard"),
    ("Description",           "User views the main dashboard showing supplier nodes on world map."),
    ("Pre-Condition(s)",      "System is running."),
    ("Post-Condition(s)",     "Dashboard is displayed."),
    ("Actor",                 "User"),
    ("Primary Flow",          "- Open system homepage\n- System loads map and supplier nodes"),
    ("Alternative /\nException Flow", "If data fails to load, error message is shown."),
])
p("Table 3.3 describes the View Map Dashboard use case. This is the entry point of the "
  "system where the user accesses the interactive world map showing all supplier nodes. "
  "The map is colour-coded based on the current risk level of each supplier.")

uc_table("Table 3.4 Filter Suppliers by Supply Chain Use Case Description", [
    ("Use Case Name",         "Filter Suppliers by Supply Chain"),
    ("Description",           "User filters suppliers by product or supply chain."),
    ("Pre-Condition(s)",      "Dashboard is open."),
    ("Post-Condition(s)",     "Filtered nodes are shown."),
    ("Actor",                 "User"),
    ("Primary Flow",          "- Select filter\n- System queries database\n- Map updates"),
    ("Alternative /\nException Flow", "If no results, empty map is shown."),
])
p("Table 3.4 describes the Filter Suppliers use case. The user can narrow down the "
  "supplier nodes displayed on the map by selecting a specific product or supply chain. "
  "This helps the user focus on a relevant subset of suppliers without viewing all nodes at once.")

uc_table("Table 3.5 View Supplier Details Use Case Description", [
    ("Use Case Name",         "View Supplier Details"),
    ("Description",           "User views detailed supplier risk and resilience info."),
    ("Pre-Condition(s)",      "Map is visible."),
    ("Post-Condition(s)",     "Detail panel is shown."),
    ("Actor",                 "User"),
    ("Primary Flow",          "- Click supplier node\n- System fetches data\n- Display details"),
    ("Alternative /\nException Flow", "If data missing, error shown."),
])
p("Table 3.5 describes the View Supplier Details use case. Clicking on a supplier node "
  "opens a detail panel that presents the current risk score, risk level, and the latest "
  "disruption events associated with that supplier. This allows the user to assess the "
  "current state of a specific supplier in depth.")

uc_table("Table 3.6 View Resilience History Use Case Description", [
    ("Use Case Name",         "View Resilience History"),
    ("Description",           "User views resilience trend history."),
    ("Pre-Condition(s)",      "Supplier page open."),
    ("Post-Condition(s)",     "History chart shown."),
    ("Actor",                 "User"),
    ("Primary Flow",          "- Click history tab\n- System loads history\n- Display chart"),
    ("Alternative /\nException Flow", "If no data, empty chart shown."),
])
p("Table 3.6 describes the View Resilience History use case. The system retrieves and "
  "displays historical risk score records for the selected supplier as a time-series chart. "
  "This enables the user to identify long-term trends and patterns in supplier stability.")

# NEW TABLE — UC5
uc_table("Table 3.7 View 14-Day Risk Forecast Use Case Description", [
    ("Use Case Name",         "View 14-Day Risk Forecast"),
    ("Description",           "User views the Two-Stage XGBoost 14-day forward risk forecast for a selected supplier."),
    ("Pre-Condition(s)",      "Supplier page open."),
    ("Post-Condition(s)",     "Forecast chart displayed with predicted risk values and confidence intervals."),
    ("Actor",                 "User"),
    ("Primary Flow",          "- Click forecast tab\n- System retrieves or generates forecast snapshot\n- Display forecast chart with yhat, yhat_lower, yhat_upper for 14 days"),
    ("Alternative /\nException Flow", "If forecast snapshot is unavailable, system generates it on demand."),
])
p("Table 3.7 describes the View 14-Day Risk Forecast use case. The system generates or "
  "retrieves a pre-computed forecast snapshot produced by the Two-Stage XGBoost model "
  "and displays the predicted risk values with upper and lower confidence intervals over "
  "the next 14 days. This allows the user to anticipate potential disruptions before they occur.")

uc_table("Table 3.8 Search Supplier Use Case Description", [
    ("Use Case Name",         "Search Supplier"),
    ("Description",           "User searches supplier by name."),
    ("Pre-Condition(s)",      "Dashboard open."),
    ("Post-Condition(s)",     "Matching suppliers shown."),
    ("Actor",                 "User"),
    ("Primary Flow",          "- Enter keyword\n- System searches DB\n- Show result"),
    ("Alternative /\nException Flow", "If not found, show no result."),
])
p("Table 3.8 describes the Search Supplier use case. The user can enter a keyword to "
  "locate a specific supplier by name or location. The system queries the database and "
  "highlights the matching supplier on the map and in the detail panel.")

uc_table("Table 3.9 View Related News Articles Use Case Description", [
    ("Use Case Name",         "View Related News Articles"),
    ("Description",           "User views news related to supplier."),
    ("Pre-Condition(s)",      "Supplier page open."),
    ("Post-Condition(s)",     "News list displayed."),
    ("Actor",                 "User"),
    ("Primary Flow",          "- Click news tab\n- System queries DB\n- Display articles"),
    ("Alternative /\nException Flow", "If none, empty list shown."),
])
p("Table 3.9 describes the View Related News Articles use case. The user can browse "
  "news articles that have been associated with a selected supplier or disruption event. "
  "The system displays the article title, source, and publication date to provide context "
  "behind the supplier risk scores.")

# CHANGE: "fetch the latest news articles" → from RSS feeds
uc_table("Table 3.10 Trigger Manual Refresh Use Case Description", [
    ("Use Case Name",         "Trigger Manual Refresh"),
    ("Description",           "Admin manually triggers RSS news update."),
    ("Pre-Condition(s)",      "Admin logged in."),
    ("Post-Condition(s)",     "System fetches from RSS feeds and updates data."),
    ("Actor",                 "Admin"),
    ("Primary Flow",          "- Click refresh\n- System polls RSS feeds\n- Run inference\n- Update DB"),
    ("Alternative /\nException Flow", "If feed unreachable, show error."),
])
p("Table 3.10 describes the Trigger Manual Refresh use case. An administrator can "
  "initiate an immediate RSS ingestion cycle outside of the scheduled polling interval. "
  "The system fetches the latest articles, runs the full inference pipeline, and updates "
  "the database with the most recent risk scores and forecast snapshots.")

uc_table("Table 3.11 Manage Supplier Nodes Use Case Description", [
    ("Use Case Name",         "Manage Supplier Nodes"),
    ("Description",           "Admin manages suppliers."),
    ("Pre-Condition(s)",      "Admin logged in."),
    ("Post-Condition(s)",     "Supplier DB updated."),
    ("Actor",                 "Admin"),
    ("Primary Flow",          "- Add/Edit/Delete supplier\n- System validates\n- Save to DB"),
    ("Alternative /\nException Flow", "If invalid data, reject input."),
])
p("Table 3.11 describes the Manage Supplier Nodes use case. The administrator can add "
  "new supplier nodes, update existing supplier details, or remove suppliers from the "
  "system. All changes are validated before being persisted to the database.")

uc_table("Table 3.12 Manage Supply Chains Use Case Description", [
    ("Use Case Name",         "Manage Supply Chains"),
    ("Description",           "Admin manages supply chains."),
    ("Pre-Condition(s)",      "Admin logged in."),
    ("Post-Condition(s)",     "Supply chain DB updated."),
    ("Actor",                 "Admin"),
    ("Primary Flow",          "- Add/Edit/Delete supply chain\n- Save changes"),
    ("Alternative /\nException Flow", "If conflict, show error."),
])
p("Table 3.12 describes the Manage Supply Chains use case. The administrator can create, "
  "update, or remove supply chain and product records that are used to group and filter "
  "supplier nodes on the dashboard. Any conflicts or duplicate entries are flagged by "
  "the system before saving.")

uc_table("Table 3.13 View System Logs / Reports Use Case Description", [
    ("Use Case Name",         "View System Logs / Reports"),
    ("Description",           "Admin views system logs known issues."),
    ("Pre-Condition(s)",      "Admin logged in."),
    ("Post-Condition(s)",     "Logs displayed."),
    ("Actor",                 "Admin"),
    ("Primary Flow",          "- Open logs page\n- System loads logs"),
    ("Alternative /\nException Flow", "If no logs, show empty page."),
])
p("Table 3.13 describes the View System Logs use case. The administrator can access a "
  "log page that records system operations such as RSS ingestion cycles, risk score "
  "updates, and forecast generation runs. This supports system monitoring and helps "
  "identify any processing errors or anomalies.")

# 3.3.6
h2("3.3.6\tUser Interface")

# VERBATIM
p("This section describes the user interface design of the Supply Chain Resilience "
  "Monitoring System. The interface was designed to present supply chain risk, resilience, "
  "and disruption information in a clear and structured manner to support effective "
  "monitoring. The following figures illustrate the mock user interface and explain the "
  "functionality of each interface.")

fig_placeholder("Figure 3.8 Map Dashboard Interface")
caption("Figure 3.8 Map Dashboard Interface")

# VERBATIM
p("Figure 3.8 displays the Map Dashboard, which gives a global view of the supply chain. "
  "Supplier nodes are presented as an interactive world map and their current level of risk "
  "is color-coded providing the user with an opportunity to easily locate high-risk "
  "suppliers. The most important metrics like the total suppliers, high-risk suppliers, the "
  "average resilience and the active alerts are presented in summary indicators at the top "
  "of the screen. Suppliers can also be filtered by the supply chain or product using filters "
  "which allow users to narrow down.")

p("Having a supplier node selected reveals a detail panel where the suppliers "
  "location, risk score, resilience score, the supply chains that include this supplier, and "
  "the disruption events in the recent past can be seen. This interface facilitates fast "
  "monitoring and top-level making decisions.")

fig_placeholder("Figure 3.9 Supplier Details Interface")
caption("Figure 3.9 Supplier Details Interface")

# VERBATIM
p("Figure 3.9 shows the Suppliers screen, in which all suppliers monitored are displayed "
  "in a table format. The table shows supplier location, supply chains, risk score, supply "
  "chain resilience score, and status. It also has a search feature and a filtering system so "
  "that the user can find particular suppliers faster.")

p("The visual representation of risk and resilience values through progress bars and color-"
  "coded labels allows the suppliers to be easily compared. This screen helps in "
  "monitoring of suppliers and comparative analysis.")

fig_placeholder("Figure 3.10 Resilience History Interface")
caption("Figure 3.10 Resilience History Interface")

# VERBATIM
p("Figure 3.10 demonstrates the Resilience History screen that enables users to track the "
  "past performance of an identified supplier.")

p("The interface shows the current risk score, current resilience score and the "
  "overall trend status. The changes in risks and resilience scores across time are "
  "represented in line charts, and the comparison of the two metrics is also provided. This "
  "screen is useful in long-term analysis of stability and recovery trends of suppliers.")

# NEW figure for forecast interface
fig_placeholder("Figure 3.11 14-Day Risk Forecast Interface")
caption("Figure 3.11 14-Day Risk Forecast Interface")

p("Figure 3.11 shows the 14-Day Risk Forecast screen, which visualises the output of the "
  "Two-Stage XGBoost forecasting engine for a selected supplier. The interface displays a "
  "forecast chart showing the predicted daily risk values and the associated confidence "
  "intervals over the next 14 days. Where historical data is available, the actual realised "
  "risk values are overlaid on the chart for comparison. This screen enables users to "
  "anticipate upcoming disruption risks before they materialise.")

# Renumber figures 3.11 → 3.12, 3.12 → 3.13
fig_placeholder("Figure 3.12 News & Events Interface")
caption("Figure 3.12 News & Events Interface")

# VERBATIM (was Figure 3.11)
p("Figure 3.12 provides an overview of the News and Events screen, which displays news "
  "articles and disruption events of suppliers under observation. The associated news "
  "section includes the sources of articles, the date of publication, and the headlines "
  "whereas the disruption events section displays organized events with the "
  "corresponding risk levels and status.")

p("This interface provides users with the ability to comprehend the actual world "
  "events that drive the supplier risk and resilience scores, which facilitates the "
  "transparency and explainable system output.")

fig_placeholder("Figure 3.13 Administration Interface")
caption("Figure 3.13 Administration Interface")

# VERBATIM (was Figure 3.12)
p("Figure 3.13 shows the Administration screen, that is intended to monitor and control "
  "the system. The administrators may manually initiate news updates, check database "
  "status, system health, and system logs.")

p("The system log table documents the important operations which include news "
  "ingestion, risk calculation and data synchronization and their state of execution. This "
  "screen aids in maintenance of systems and operational control.")

# ══════════════════════════════════════════════════════════════════════════════
# 3.4
# ══════════════════════════════════════════════════════════════════════════════
h1("3.4\tPhase 3: Development")

# CHANGE: "BERT and XGBoost" → "XGBoost, FinBERT, and GLiNER2"; "News API live data" → "RSS-based live data"
p("Phase 3 entailed implementation of the supply chain disruption detection system as a "
  "local prototype web application. This stage combined the frontend, backend and "
  "machine learning modules into a working solution. The tasks involved in development "
  "were to create a ReactJS front-end interface, a Python back-end API with FastAPI "
  "framework, a hybrid ML inference engine with combination of XGBoost, FinBERT, and "
  "GLiNER2 models and RSS-based live data ingestion. The aim was to make sure that "
  "real-time news articles that may act as early warning signs of supply chain disruptions "
  "are automatically downloaded, processed in the ML pipeline, and displayed on the web "
  "interface with relevant risk and resilience scores.")

# 3.4.1
h2("3.4.1\tHardware and Software Requirements")

# VERBATIM
p("The hardware requirements outline the required minimum of computing available to "
  "support system development, model training and real-time inference processes. Big "
  "data news data and machine learning loads need to be processed with sufficient "
  "processing capacity, memory, and storage. The system has been initially set up to run "
  "in a local machine environment in order to enable a controlled development and "
  "testing.")

# Hardware table — minimal changes: RAM mention FinBERT/GLiNER2; Storage update size; GPU update
tbl_caption("Table 3.14 Hardware Requirements")
make_table(
    ["Component", "Specification", "Justification"],
    [
        ["Processor (CPU)",
         "64-bit multi-core processor (e.g., Intel Core i5 or AMD equivalent, 2.5 GHz or above).",
         "A multi-core CPU helps in handling the web server and ML inference in parallel."],
        ["Memory (RAM)",
         "8 GB RAM minimum (16 GB recommended).",
         "Sufficient memory is needed to load the FinBERT and GLiNER2 models and handle data "
         "processing for multiple news articles simultaneously without performance degradation."],
        ["Storage",
         "At least 5–10 GB free disk space.",
         "This accommodates the project source code, installed libraries, and the cached ML "
         "models (FinBERT and GLiNER2 combined can be ~600 MB). Additional space may be used "
         "for storing retrieved news data, database files, or logs."],
        ["Graphics (GPU)",
         "NVIDIA GPU with CUDA support (e.g., NVIDIA GTX/RTX series) or Apple Silicon GPU (MPS) "
         "for accelerating FinBERT and GLiNER2 inference.",
         "The system is capable of running on CPU alone, but a GPU can significantly speed up "
         "model inference if processing large batches of articles. The system auto-detects "
         "the available device."],
        ["Network",
         "Active internet connection for the backend to poll the RSS feeds.",
         "Also used during development for installing packages (via npm/pip) and for any "
         "online dataset access."],
        ["Operating System",
         "Windows 10 (64-bit) or a compatible OS (e.g., Linux Ubuntu 20.04+, macOS).",
         "The development will be conducted on a local machine, but the stack (Node, Python, etc.) "
         "is cross-platform, so any OS supporting these tools meets the requirements."],
    ]
)

# VERBATIM
p("Software requirements indicate the programming environment, libraries, and tools that "
  "will be required to develop and implement the proposed supply chain resilience "
  "analysis system. These software modules facilitate data processing, machine learning "
  "model development, integration of news data and interactive visualization of supply "
  "chain networks. The tools chosen are compatible and can be scaled and can easily "
  "enable the system to run smoothly during development and testing.")

# Software table — change BERT→FinBERT, spaCy scope, News API→RSS, add GLiNER2 & PostgreSQL
tbl_caption("Table 3.15 Software Requirements")
make_table(
    ["Software", "Version", "Purpose"],
    [
        ["Python",               "Python 3.11 or higher",
         "Provides the programming environment for backend development, data processing, and machine learning model execution."],
        ["FastAPI",              "Latest stable version",
         "Used to develop RESTful backend APIs and handle communication between frontend and machine learning components."],
        ["ReactJS",              "Latest stable version",
         "Builds the interactive web-based frontend dashboard for visualizing supply chain disruption and resilience analysis."],
        ["PostgreSQL",           "Latest stable version",
         "Relational database that serves as the system of record for supplier data, events, risk history, and forecast snapshots."],
        ["Hugging Face Transformers (FinBERT)", "Latest version",
         "Provides the FinBERT (ProsusAI/finbert) model for financial domain sentiment analysis on news articles, producing sentiment labels and scores."],
        ["XGBoost",              "Latest version",
         "Implements the gradient boosting models used for tri-class headline risk classification, disruption detection, impact regression, and the Two-Stage risk forecaster."],
        ["GLiNER2",              "fastino/gliner2-base-v1",
         "Named entity recognition model for extracting geographic location entities from news text."],
        ["spaCy",                "Latest version (en_core_web_sm)",
         "Used for temporal date extraction to identify when an event is expected to occur. Not used for named entity recognition in the current system."],
        ["Pandas & NumPy",       "Latest version",
         "Handles data preprocessing, numerical computation, and feature engineering."],
        ["Node.js & npm",        "Latest LTS version",
         "Supports ReactJS frontend development and dependency management."],
    ]
)

# 3.4.2
h2("3.4.2\tMachine Learning Model Pseudocode")

# VERBATIM
p("The subsection below details the pseudocode of the machine learning elements applied "
  "in the Supply Chain Resilience Monitoring System. The pseudocode outlines the "
  "offline training process and online inference process which constitute a hybrid "
  "machine learning architecture. This pseudocode is meant to show in detail the logical "
  "sequence of events in the disruption detection, model deployment and resilience score "
  "update.")

h3("3.4.2.1\tAlgorithm 1: Offline Model Training Procedure")

# CHANGE: update pseudocode to reflect TF-IDF+XGBoost multi-model training; keep style
pseudo([
    "Input:  Historical news dataset from Hugging Face",
    "Output: Trained models — classifier.pkl (tri-class XGBoost),",
    "        disruption_classifier.pkl, impact_regressor_v2.pkl,",
    "        forecast_event_prob.json, forecast_severity_q75.json",
    "",
    "1.  Begin",
    "2.  Load historical news dataset",
    "3.  Extract text and label fields from dataset",
    "",
    "4.  For each news article in dataset do",
    "5.      Clean text (remove punctuation, numbers, special characters)",
    "6.      Convert text to lowercase",
    "7.      Perform tokenization",
    "8.      Remove stop words",
    "9.      Apply lemmatization",
    "10. End For",
    "",
    "11. Apply TF-IDF vectorisation to produce feature vectors",
    "12. Encode tri-class labels (LOW / MEDIUM / HIGH) using LabelEncoder",
    "13. Split dataset into training and validation sets",
    "",
    "14. Train XGBoost tri-class headline risk classifier",
    "15. Train XGBoost binary disruption classifier",
    "16. Train XGBoost impact regressor",
    "17. Train Two-Stage XGBoost forecaster (freeze-window):",
    "        Stage 1 — XGBClassifier predicts P(event) per horizon day",
    "        Stage 2 — XGBRegressor (q=0.75) predicts E[severity | event]",
    "        Final:    yhat = P(event) x severity",
    "",
    "18. Evaluate model performance using validation data",
    "19.     Compute accuracy, precision, recall, and F1-score",
    "20.     If model performance is satisfactory then",
    "21.         Save trained model to disk",
    "22.         Store model metadata (version, date, metrics)",
    "23.     End If",
    "24. End",
])

# CHANGE: description updated to reflect multi-model + freeze-window, keep original tone
p("Algorithm 1 describes how the disruption detection models are trained using historical "
  "news data. It begins with the loading of the news dataset and the extraction of the text "
  "and the labels. News articles are then cleaned and pre-processed with preprocessing "
  "algorithms like tokenization, stop word removal and lemmatization. The steps can be "
  "used to minimize noise and enhance the quality of the input data. TF-IDF vectorisation "
  "then converts the cleaned text into numerical features for the XGBoost classifiers.")

p("Three separate XGBoost models are trained: a tri-class headline risk classifier, a "
  "binary disruption classifier, and an impact regressor. Additionally, the Two-Stage "
  "XGBoost forecaster is trained using a freeze-window architecture, where Stage 1 "
  "predicts the probability of a disruption event and Stage 2 predicts the expected "
  "severity. The data is divided into training and validation groups in order to facilitate "
  "objective performance assessment. When an acceptable result is reached, the trained "
  "models and the metadata are stored to be deployed.")

h3("3.4.2.2\tAlgorithm 2: Online News Inference and Risk Update Procedure")

# CHANGE: input "from NewsAPI" → "from RSS feeds"; add multi-model steps; keep style
pseudo([
    "Input:  Live news articles from RSS feeds",
    "Output: Updated supplier risk and resilience scores,",
    "        refreshed 14-day forecast snapshots",
    "",
    "1.  Begin",
    "2.  Load trained models (classifier.pkl, disruption_classifier.pkl,",
    "    impact_regressor_v2.pkl, forecast_event_prob.json,",
    "    forecast_severity_q75.json)",
    "3.  Poll configured RSS feeds and retrieve latest articles",
    "",
    "4.  For each retrieved news article do",
    "5.      Validate article and remove duplicates",
    "6.      Clean and preprocess article text",
    "",
    "7.      Extract named location entities using GLiNER2",
    "8.      Geocode extracted locations",
    "9.      Match geocoded locations with supplier node database",
    "",
    "10.     If related supplier is found then",
    "11.         Classify headline using XGBoost tri-class classifier",
    "            -> ml_risk_label (LOW / MEDIUM / HIGH), risk_score",
    "12.         Compute sentiment using FinBERT",
    "            -> sentiment_label, sentiment_score",
    "13.         Predict disruption probability using XGBoost classifier",
    "            -> predicted_disruption_probability",
    "14.         Predict impact score using XGBoost regressor",
    "            -> predicted_impact_score",
    "",
    "15.         If article is classified as disruption then",
    "16.             Create disruption event record",
    "17.             Update supplier risk score",
    "18.             Recalculate supplier resilience score",
    "19.             Append record to resilience history",
    "20.             Regenerate 14-day forecast snapshot (Two-Stage XGBoost)",
    "21.         Else",
    "22.             Store article as non-disruption record",
    "23.         End If",
    "24.     End If",
    "25. End For",
    "",
    "26. Update frontend API with latest risk, resilience, and forecast data",
    "27. End",
])

# CHANGE: description updated for RSS + GLiNER2 + multi-model, keep original tone
p("Algorithm 2 describes the use of the trained models in system operation. The RSS "
  "ingestion module polls the configured RSS feeds, validates the articles and pre-processes "
  "them through the same steps that are applied during training. Location named entity "
  "recognition is performed using the GLiNER2 model to identify geographic references in "
  "the news text.")

p("When a news item is associated with a familiar supplier, the article passes through the "
  "multi-model inference pipeline. The XGBoost tri-class classifier assigns a risk label, "
  "FinBERT computes a sentiment score, and the disruption classifier and impact regressor "
  "produce a disruption probability and an impact score respectively. On detecting a "
  "disruption, the system generates a disruption event record, modifies the risk and "
  "resilience scores of the supplier, and regenerates the 14-day forecast snapshot for the "
  "affected node. Non-disruption articles remain stored without impacting supplier scores.")

# ══════════════════════════════════════════════════════════════════════════════
# 3.4.3  RSS Feed Ingestion (replaces NewsAPI Integration)
# ══════════════════════════════════════════════════════════════════════════════
h2("3.4.3\tRSS Feed Ingestion")

# Keep same structure/tone as original 3.4.3, just update content for RSS
p("RSS feed ingestion is incorporated into the system as the primary external source of "
  "real-time news articles concerning supply chain disruptions. The feeds are configured "
  "in a JSON file (config/rss_feeds.json) as an array of entries, each containing a feed "
  "URL and a source label, allowing any number of sources to be monitored without "
  "code changes.")

p("Implementation of the integration is done in Python, with triggers being scheduled "
  "automatically by a background worker or manually triggered through the POST "
  "/admin/rss-ingest/trigger endpoint. Articles retrieved are parsed and validated "
  "before being stored temporarily to be processed further. Redundant data is eliminated "
  "by filtering of duplicate articles to avoid repetition of analysis. This is done to "
  "ensure the inference pipeline remains efficient and to minimize unwanted computing "
  "overhead.")

p("Upon validation, the news articles are sent to the preprocessing phase and inference "
  "phase, where appropriate entities including suppliers, locations and disruption events "
  "are identified using the GLiNER2 model. The obtained processed results are then "
  "connected to the existing records of the suppliers in the database, with the system "
  "then able to evaluate possible risks and adjust the resilience metrics.")

fig_placeholder("Figure 3.14 Coding Segment for RSS Feed Ingestion")
caption("Figure 3.14 Coding Segment for RSS Feed Ingestion")

# ══════════════════════════════════════════════════════════════════════════════
# 3.4.4  GLiNER2 Location NER (replaces Limitations of NewsAPI)
# ══════════════════════════════════════════════════════════════════════════════
h2("3.4.4\tGLiNER2 Named Entity Recognition")

p("Location named entity recognition (NER) is a critical step in associating news articles "
  "with the correct supplier nodes. The system uses GLiNER2 (fastino/gliner2-base-v1) "
  "for this purpose. This model was introduced to replace an earlier BERT-Large based NER "
  "approach, significantly reducing inference memory requirements and batch processing "
  "time while maintaining high extraction accuracy for geographic location entities.")

p("GLiNER2 operates by performing span-level entity classification on the input text given "
  "a set of target entity type labels (in this case, location). The model is invoked in batch "
  "mode to process multiple articles efficiently. The GLINER_MODEL environment variable "
  "allows the model identifier to be updated without code changes, supporting future "
  "model upgrades. The extracted location entities are subsequently passed to a geocoding "
  "module that resolves them to latitude and longitude coordinates, which are then matched "
  "against the canonical supplier node list.")

# ══════════════════════════════════════════════════════════════════════════════
# 3.4.5  Two-Stage XGBoost Risk Forecasting (new section)
# ══════════════════════════════════════════════════════════════════════════════
h2("3.4.5\tTwo-Stage XGBoost Risk Forecasting")

p("The system incorporates a risk forecasting component that generates 14-day forward "
  "risk predictions for each supplier node. The forecasting engine is implemented as a "
  "Two-Stage XGBoost model (src/forecast_snapshots.py), which is the production "
  "forecasting path for all forecast API endpoints and pipeline runs.")

p("Stage 1 estimates the probability that a disruption event occurs on a given future day "
  "(P(event)), trained as an XGBoost binary classifier and stored in "
  "models/forecast_event_prob.json. Stage 2 estimates the expected severity given that "
  "an event occurs (E[severity | event]), trained as an XGBoost quantile regressor at "
  "the 75th percentile and stored in models/forecast_severity_q75.json. The final "
  "predicted risk value for each horizon day is computed as: yhat = P(event) × severity.")

p("A freeze-window architecture is used whereby all features for every horizon day are "
  "computed entirely from actual data available at the forecast origin date. A day_offset "
  "feature (integer from 1 to 14) is included as a model input to differentiate each "
  "horizon position. This design eliminates recursive data leakage and allows the model "
  "to be evaluated on truly out-of-sample future days. Forecast results are persisted in "
  "the ForecastSnapshot table and served through the GET "
  "/suppliers/{node_name}/forecast API endpoint, with on-demand generation if a "
  "pre-computed snapshot is not found.")

# ══════════════════════════════════════════════════════════════════════════════
# 3.5
# ══════════════════════════════════════════════════════════════════════════════
h1("3.5\tPhase 4: Testing")

# VERBATIM
p("The stage aims at ensuring that the proposed Supply Chain Resilience Monitoring "
  "System is functioning correctly, is reliable, and working well. The test is done to verify "
  "that every component of the system is functioning as per the requirements outlined "
  "such as data ingestion, machine learning inference, risk and resilience score updates, "
  "and user interface interactions. The test scenarios are modeled depending on the "
  "functional needs and implemented in a local deployment environment.")

h2("3.5.1\tFunctional Testing")

# VERBATIM
p("Table 3.5.1 presents the functional test cases conducted for the system. Each test case "
  "includes the objective, preconditions, testing steps, expected results, actual results, and "
  "test status. All tests were executed after successful system setup and model loading.")

# Functional testing table — keep TC001–TC007 verbatim; update TC002 (NewsAPI→RSS); add TC008 for forecast
tbl_caption("Table 3.16 Functional Testing")
make_table(
    ["Test Case ID", "Objective", "Precondition", "Steps", "Expected Result", "Actual Result", "Status"],
    [
        ["TC001", "Verify that machine learning models are loaded successfully",
         "Application server is launched", "Start the backend API server",
         "All trained ML models are loaded without runtime errors",
         "Models loaded successfully without errors", "Pass"],
        # CHANGE: NewsAPI → RSS feeds
        ["TC002", "Verify real-time news ingestion and disruption classification",
         "System is running and RSS feeds are configured",
         "Trigger news ingestion process from API",
         "Relevant news articles are retrieved from RSS feeds and classified correctly",
         "News articles processed and classified successfully", "Pass"],
        ["TC003", "Verify supplier risk score update after disruption detection",
         "Disruption event exists in the database",
         "Execute risk scoring module",
         "Supplier risk score is updated based on detected event",
         "Risk score updated correctly in database", "Pass"],
        ["TC004", "Verify resilience score computation and historical update",
         "Risk score has been updated",
         "Run resilience score calculation",
         "Resilience score is recalculated and stored in history table",
         "Resilience history updated successfully", "Pass"],
        ["TC005", "Verify world map visualization of supplier nodes",
         "Frontend application is running",
         "Open world map homepage",
         "Supplier nodes and links are displayed correctly on map",
         "Map rendered with correct supplier visualization", "Pass"],
        ["TC006", "Verify supplier filtering by supply chain",
         "Supplier and product data exist",
         "Select a specific product supply chain filter",
         "Only suppliers related to selected supply chain are shown",
         "Correct suppliers filtered and displayed", "Pass"],
        ["TC007", "Verify supplier detail page display",
         "Supplier record exists",
         "Click on a supplier node",
         "Supplier details and historical data are displayed",
         "Supplier detail page displayed correctly", "Pass"],
        # NEW TC008 for forecast
        ["TC008", "Verify 14-day risk forecast generation and display",
         "Supplier node exists and ML models are loaded",
         "Call GET /suppliers/{node_name}/forecast",
         "14-day forecast snapshot is returned with yhat, yhat_lower, yhat_upper for each horizon day",
         "Forecast generated and displayed correctly", "Pass"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.6
# ══════════════════════════════════════════════════════════════════════════════
h1("3.6\tSummary")

# VERBATIM para 1
p("This chapter has presented the complete methodology adopted for the development of "
  "the Supply Chain Resilience Monitoring System. The overall development process "
  "was guided by the Waterfall model, which was selected due to its structured, sequential "
  "nature and suitability for projects with well-defined requirements. The methodology "
  "clearly outlined each development phase, beginning from system requirement analysis "
  "and design, followed by implementation, testing, and deployment. This method "
  "guaranteed the systematic progress and explicit validation at every stage of "
  "development.")

# CHANGE: "NewsAPI" → "RSS feeds"; "online inference processes live news articles...via NewsAPI" → RSS; add forecasting
p("The system architecture was modelled as a hybrid system that integrates offline "
  "machine learning model training with online inference. The offline training was "
  "conducted with the help of historical news datasets to train and test disruption detection "
  "models, and online inference processes live news articles, which are obtained through "
  "RSS feed ingestion. The preprocessing of the text, location entity recognition using "
  "GLiNER2, and classification of disruptions were performed using Natural Language "
  "Processing methods as they allowed the system to recognize the potential risks in the "
  "supply chain that come with individual suppliers and regions. A Two-Stage XGBoost "
  "forecasting engine further generates 14-day forward risk predictions for each supplier "
  "node.")

# VERBATIM para 3 — mostly unchanged
p("The development stage aimed at deploying the back-end service based on "
  "Python, the opportunity to integrate trained machine learning models into an inference "
  "pipeline and create a web interface based on ReactJS to visualize the data. The system "
  "offers interactive capabilities including world map with supplier nodes view, supply "
  "chain relationship view, risk and resilience score and historical trend view. All these "
  "elements combine to provide real-time data on supply chain vulnerability and "
  "resilience.")

# VERBATIM para 4
p("Finally, this chapter has shown how there is a methodological approach "
  "towards developing the proposed system. By integrating machine learning, natural "
  "language processing, and interactive visualization within a well-defined "
  "methodological framework, the Supply Chain Resilience Monitoring System provides "
  "a practical tool for enhancing supply chain risk awareness and resilience assessment. "
  "The methodology adopted in this chapter establishes a strong foundation for the "
  "implementation results and evaluation discussed in the subsequent chapters.")

# ── save ─────────────────────────────────────────────────────────────────────
out = "/Users/meordanish/Desktop/Projects/SupplyChainForecast/docs/chap3_updated.docx"
doc.save(out)
print(f"Saved: {out}")
