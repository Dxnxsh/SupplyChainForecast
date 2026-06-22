"""
Generate updated Chapter 3 docx for FYP submission.
Run: venv311/bin/python scripts/generate_chap3_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.0)

# ── helper: paragraph style ──────────────────────────────────────────────────
def body(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(24)   # double-ish
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.25)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p

def fig_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    return p

def table_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p

def placeholder_figure(label):
    """Light-grey placeholder box where a figure will be inserted."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(5.5)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "D9D9D9")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after  = Pt(40)
    run = p.add_run(f"[{label}]")
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()   # spacing after

def add_table(headers, rows, caption):
    table_caption(caption)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for r_idx, row in enumerate(rows):
        tr = t.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
    doc.add_paragraph()

def pseudocode_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE (decorative)
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UNIVERSITI TEKNOLOGI MARA")
run.bold = True
run.font.size = Pt(14)
run.font.name = "Times New Roman"

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EARLY WARNING SYSTEM FOR\nSUPPLY CHAIN RESILIENCE")
run.bold = True
run.font.size = Pt(14)
run.font.name = "Times New Roman"

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MEOR DANISH FARHAN BIN SOBRI")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Thesis submitted in partial fulfilment of the requirements for\nBachelor of Computer Science (Hons.)")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Faculty of Computer and Mathematical Science")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER HEADING
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CHAPTER 3\nMETHODOLOGY")
run.bold = True
run.font.size = Pt(14)
run.font.name = "Times New Roman"
p.paragraph_format.space_after = Pt(12)

body(
    "This chapter outlines the methodology adopted for the development of the proposed system. "
    "It provides a structured methodology that supports the process of planning, designing, "
    "implementing, and evaluating the various activities across the project lifecycle. The "
    "software development model that will be used in this study is the Waterfall Model that is "
    "defined by the sequential and linear flow of development activities."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.1  Software Development Methodology
# ══════════════════════════════════════════════════════════════════════════════
heading1("3.1\tSoftware Development Methodology")

body(
    "The waterfall model is a software development methodology that is based on the software "
    "development lifecycle and employs a sequential and linear approach. It is used for projects "
    "that require clarity, consistency, and are unlikely to change while they are being developed. "
    "The waterfall model structure is shown in greater depth in Figure 3.1."
)

placeholder_figure("Figure 3.1 – Insert Waterfall Model Structure diagram")
fig_caption("Figure 3.1 Waterfall Model Structure")

add_table(
    ["Objective", "Phase", "Activities", "Outcome"],
    [
        [
            "Objective 1: To identify suitable techniques for monitoring and assessing supply chain disruption risks.",
            "Planning",
            "• Conduct literature review on supply chain disruption, resilience, and risk assessment techniques.\n"
            "• Review existing machine learning and NLP techniques used for disruption detection.\n"
            "• Analyse suitable datasets and data sources for supply chain risk monitoring.",
            "• Selection of appropriate machine learning and NLP techniques.\n"
            "• Identification of relevant datasets and disruption indicators."
        ],
        [
            "Objective 2: To design a supply chain resilience monitoring system using a hybrid machine learning architecture.",
            "Design\nDevelopment",
            "• Design overall system architecture and workflow diagrams.\n"
            "• Design offline training and online inference pipelines.\n"
            "• Design database schema and user interface structure.\n"
            "• Implement backend services for data ingestion and processing.\n"
            "• Train machine learning models using historical news data.\n"
            "• Integrate trained models into the inference pipeline.\n"
            "• Develop frontend dashboard for visualization.",
            "• Defined system architecture and data flow.\n"
            "• Finalised ERD and interface design.\n"
            "• Trained and stored machine learning models.\n"
            "• Functional web-based supply chain monitoring system."
        ],
        [
            "Objective 3: To test the functionality and reliability of the developed supply chain resilience monitoring system.",
            "Testing",
            "• Perform functional testing on system modules.\n"
            "• Validate model inference and score updates.\n"
            "• Test user interface and visualization components.",
            "• Verified system functionality and performance.\n"
            "• Identified potential improvements."
        ],
    ],
    "Table 3.1 Project Framework"
)

body(
    "Table 3.1 presents the project framework that maps each project objective to the "
    "corresponding phase of the Waterfall model. Objective 1 is addressed during the planning "
    "phase through literature review and technique selection. Objective 2 is achieved in the "
    "design and development phases through system design, model training, and system "
    "implementation. Lastly, Objective 3 is fulfilled in the testing phase by evaluating system "
    "functionality and accuracy performance."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.2  Phase 1: Requirement Analysis
# ══════════════════════════════════════════════════════════════════════════════
heading1("3.2\tPhase 1: Requirement Analysis")

body(
    "During this stage of planning, the critical goals of the project, data needs, and approaches "
    "to the methodology are established to aid in the successful implementation of the supply "
    "chain disruption forecasting system. It is centred on how the textual data obtained from news "
    "articles could be used to identify and forecast events that are likely to interfere with the "
    "operations in the supply chain. This involves determining the features that are relevant to "
    "news content, including the type of event (e.g., strikes, disasters, geopolitical problems), "
    "entities (places, organisations), and time."
)

body(
    "In that regard, the project uses the News Dataset offered by Hugging Face "
    "(R3troR0b/news-dataset), a collection of real-life news headlines and article summaries "
    "that were annotated with event categories. In contrast to other types of disruption "
    "detection models, which utilise social media or financial information, this system was "
    "created to be trained exclusively on news-based text information to guarantee the "
    "reliability of content and richness of context."
)

body(
    "At this stage, suitable machine learning models are also selected, especially supervised "
    "classifiers such as XGBoost, FinBERT, and GLiNER2, depending on past performance "
    "comparisons. The planning stage also involves setting up the software tools and libraries "
    "to work with (e.g. Python, scikit-learn, Hugging Face Transformers, XGBoost), system "
    "objectives, and aligning the data pipeline with the system architecture in general."
)

# 3.2.1
heading2("3.2.1\tData Collection")

body(
    "The main dataset that is used in this study is the R3troR0b/news-dataset that can be found "
    "in Hugging Face. This dataset consists of a large amount of news articles from across the "
    "globe. Based on its dataset card, the articles are pulled from reputable publishers such as "
    "Reuters, BBC World News, Al Jazeera, Le Monde, South China Morning Post, The Hindu, "
    "Deutsche Welle, The Guardian, NPR, TASS (Russia), and The Sydney Morning Herald. The "
    "dataset is labelled as news and world and it is intended for text classification. It is "
    "primarily in English (with some parts in French and Russian) and is covered by the MIT "
    "Licence. The MIT licence is free to use and distribute, making it appropriate for research."
)

body(
    "The data is presented in JSON format with two primary fields per record: label (string) "
    "and text (string). The label field is a composite key that contains metadata, namely the "
    "news source (publisher), article title, URL, and the publication time "
    "(e.g. The Guardian; Article Title; https://...;2025-01-01T14:00:00Z). The news article "
    "body is contained in the text field. The dataset is updated actively: files are sorted by "
    "quarter (e.g., allnewsq12025.json, allnewsq22025.json, etc.) and are periodically updated "
    "to add new articles."
)

body(
    "The repository contains multi-megabyte JSON files per quarter, which provides a strong "
    "amount of data to train the models. Its applicability to supply chains is anticipated since "
    "global news normally reports on events which impact logistics, trade, and supply chain "
    "disruptions. Moreover, the range of the dataset (international sources and up-to-date "
    "content) includes references to keywords related to the supply chain (e.g. 'supply chain', "
    "'logistics', 'shipment', 'inventory', 'transportation' etc.). This renders it an effective "
    "foundation to develop a news classification platform centred on supply chain issues."
)

placeholder_figure("Figure 3.2 – Insert News Dataset screenshot")
fig_caption("Figure 3.2 News Dataset")

add_table(
    ["Index", "Feature Name", "Feature Explanation"],
    [
        ["F1", "Label", "This column includes the article's news provider, title, and URL link. The three data fields are separated by a semicolon (;)."],
        ["F2", "Text",  "This column includes the full-text news article."],
    ],
    "Table 3.2 Explanations for News Dataset Features"
)

# 3.2.2
heading2("3.2.2\tData Pre-Processing")

placeholder_figure("Figure 3.3 – Insert Flowchart of Data Pre-Processing")
fig_caption("Figure 3.3 Flowchart of Data Pre-Processing")

body(
    "Before classification, raw news text has to be cleaned and prepared. During this step, "
    "NLP tools (i.e., pandas, NLTK, and spaCy) are used to preprocess the articles. To work "
    "with tabular data, a JSON data block may be loaded into a pandas DataFrame using "
    "pandas.read_json(). The content of each article is extracted from the text field. The most "
    "common steps of cleaning the text are to delete punctuation, HTML tags (if they exist), and "
    "numeric characters, and to convert all text to lowercase so that there is no discrepancy."
)

body(
    "Then, tokenization takes place: every article is divided into individual words or tokens. "
    "Efficient tokenizers are offered in libraries such as spaCy and NLTK. Once tokenized, stop "
    "word removal and lemmatization are done to normalise the tokens. Stop words (such frequent "
    "words as and, the, of, etc.) bear minimal semantic information and can be filtered with the "
    "help of in-built lists in NLTK (nltk.corpus.stopwords) or spaCy (token.is_stop). "
    "Lemmatization decomposes words down to their base or dictionary form (e.g. ships → ship, "
    "running → run), which assists in grouping similar forms of words."
)

body(
    "Also, since this system is geared towards news related to supply chains, keyword filtering "
    "is applied to narrow down to relevant news. Once preprocessing is done, all the articles "
    "may be searched with key terms connected with supply chains (such as supply chain, "
    "logistics, manufacturing, procurement, transport, warehouse, etc.). Articles that contain "
    "no such keywords may be de-prioritised or classified as probably not related to supply chain "
    "classification. This filtering may be applied using simple list matching in pandas or using "
    "the phrase matcher in spaCy."
)

body(
    "It is important to note that spaCy is used exclusively in this system for temporal date "
    "extraction (i.e., identifying when an event is expected to occur) via the en_core_web_sm "
    "model. Named entity recognition (NER) for locations is handled by a separate, more "
    "specialised model described in Section 3.4.2."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.3  Phase 2: Design
# ══════════════════════════════════════════════════════════════════════════════
heading1("3.3\tPhase 2: Design")

body(
    "This part describes the general architecture of the supply chain disruption detection "
    "platform. It addresses the high-level system architecture, the data flow between the input "
    "(news articles) and the output (risk visualisation), machine learning algorithms and feature "
    "engineering, the database design, and the user interface. The main design features are a "
    "multi-layer architecture that combines data ingestion, multi-model machine learning "
    "inference, 14-day risk forecasting, and a user-interactive map frontend. The design focuses "
    "on identifying disruption events caused by unstructured news data in time and making "
    "intuitive visualisation of supplier risks, resilience, and forward-looking forecasts."
)

# 3.3.1
heading2("3.3.1\tSystem Architecture")

placeholder_figure("Figure 3.4 – Insert updated System Architecture Diagram")
fig_caption("Figure 3.4 System Architecture Diagram")

body(
    "According to Figure 3.4 (System Architecture Diagram), the system starts as the user "
    "enters the web application using the ReactJS interface. The user will be introduced to the "
    "homepage, whereby there will be a world map that contains supplier nodes denoting the "
    "various suppliers within the global supply chain. The system tracks six canonical supplier "
    "nodes: TSMC (Hsinchu), Foxconn (Zhengzhou), Port of Long Beach, CATL (Ningde), "
    "Albemarle (Chile), and Tesla (Berlin). Each node is assigned a criticality score between "
    "1 and 5 that reflects its strategic importance to the supply chain and is used to weight "
    "the final risk impact score."
)

body(
    "The system allows the user to engage with it by filtering through supplier information and "
    "viewing past resilience performance of each supplier. Whenever the user takes any action, "
    "such as filtering suppliers or choosing a particular supplier node, the request is dispatched "
    "to the backend API server. The backend processes the request using the filtering and query "
    "logic module, which retrieves applicable data in the database including supplier information, "
    "the current risk level, and historical resilience records. The information recovered is then "
    "sent back to the frontend and shown to the user as updated map nodes, detailed panels, or "
    "historical charts."
)

body(
    "The RSS ingestion scheduler continuously retrieves new articles from configured RSS feeds "
    "in the background. The fetched news items are passed to the inference orchestrator which "
    "directs the prediction process. The first step involves a preprocessing stage to clean and "
    "normalise the text. Location entities are then extracted using the GLiNER2 model. The "
    "extracted entities are geocoded and matched against the supplier database to identify "
    "the relevant supplier node. The trained multi-model ML pipeline then examines the "
    "processed text to produce a risk label, sentiment score, disruption probability, and "
    "impact score."
)

body(
    "In the case of a disruption, the system updates the risk and resilience score of the "
    "impacted supplier and saves the article and classification outcome in the database. "
    "Additionally, the Two-Stage XGBoost forecasting engine generates a 14-day forward "
    "risk forecast for the affected supplier node, which is stored in the forecast_snapshots "
    "table and surfaced through the forecast endpoint. In the case of a non-disruption article, "
    "it is still saved in the database for record-keeping and future reference."
)

body(
    "To prepare the models, the system relies on a separate offline training flow. The machine "
    "learning models are trained and tested on the Hugging Face news dataset and historical "
    "event data. When the models have sufficiently good performance, they are stored and "
    "loaded into the live inference pipeline to support real-time prediction."
)

# 3.3.2
heading2("3.3.2\tSystem Flowchart")

placeholder_figure("Figure 3.5 – Insert updated System Flowchart")
fig_caption("Figure 3.5 System Flowchart")

body(
    "According to Figure 3.5 (System Flowchart), the system process initiates when the system "
    "is triggered automatically by the scheduled RSS ingestion process or manually by an "
    "administrator. After the process begins, the system retrieves the most recent news articles "
    "from the configured RSS feeds."
)

body(
    "Following the retrieval of the news articles, the system undertakes the process of "
    "validating and deduplicating articles to ensure that only relevant articles that have never "
    "been processed before are taken into account. The articles that pass the validation process "
    "are then sent to the text preprocessing phase where the text is cleaned, tokenized, and "
    "normalised to allow it to be analysed."
)

body(
    "The system then carries out location entity extraction using the GLiNER2 model, which "
    "identifies place names and geographic references in the news text. The extracted locations "
    "are geocoded and compared with the supplier database to ascertain the relevant supplier "
    "nodes associated with the news article."
)

body(
    "After identifying the supplier association, the system runs the multi-model ML inference "
    "pipeline. The XGBoost headline classifier first assigns a tri-class risk label (LOW, MEDIUM, "
    "or HIGH) to the article. FinBERT is then applied to compute a sentiment score. The "
    "disruption classifier and impact regressor produce a disruption probability and an impact "
    "score respectively. These outputs together determine whether the article represents a "
    "supply chain disruption event."
)

body(
    "When the article is defined as a non-disruption event, the system stores the article as a "
    "regular news record, updates the news archive in the database, and ensures the most recent "
    "data is accessible to the frontend API. The frontend then updates the world map interface "
    "to show the most recent information."
)

body(
    "In the case of a disruption event, the system generates a disruption event record and "
    "associates it with the affected supplier. The system then recalculates the supplier risk "
    "score and appends the new risk information to the resilience history. Following the risk "
    "score update, the Two-Stage XGBoost forecasting engine regenerates the 14-day forward "
    "risk forecast for the node. The article, event record, and updated forecast are then "
    "persisted into the database and the frontend is updated with the latest risk and forecast "
    "data on the world map."
)

# 3.3.3
heading2("3.3.3\tEntity Relationship Diagram")

placeholder_figure("Figure 3.6 – Insert updated Entity Relationship Diagram")
fig_caption("Figure 3.6 Entity Relationship Diagram")

body(
    "According to Figure 3.6, the Entity Relationship Diagram (ERD) shows how the objects "
    "of the system database are organised and how they relate to one another. The ERD is "
    "developed to accommodate storing supply chain data, supplier details, news articles, "
    "disruption events, resilience history, and risk forecast snapshots."
)

body(
    "The Supplier table stores the details of each supplier in the form of nodes on the world "
    "map interface. Each supplier is identified with a unique identifier and attributes that "
    "include supplier name, country, region, geographic coordinates (latitude and longitude), "
    "industry, criticality score (an integer from 1 to 5 reflecting strategic importance), "
    "and an optional list of products. This table also contains the current risk score for "
    "every supplier, which is represented as an exposure index (0–100) and is regularly "
    "recomputed through the system."
)

body(
    "The SupplyChain table contains data on various product supply chains tracked by the "
    "system. The association between SupplyChain and Supplier is managed through a junction "
    "table named SupplyChainSupplier, which implements the many-to-many relationship and "
    "enables the system to filter suppliers depending on the chosen product or supply chain."
)

body(
    "The NewsArticle table contains all news articles retrieved through the RSS ingestion "
    "pipeline. Articles are uniquely identified and contain attributes including title, origin, "
    "URL, date of publication, content, and time of ingestion."
)

body(
    "The Event table stores structured disruption events extracted from news articles. Each "
    "event record contains a unique identifier, a foreign key to the NewsArticle table, event "
    "type, extracted locations (stored as a JSONB array), matched supplier nodes (stored as a "
    "JSONB array, as a single event may affect multiple nodes), risk score, relevance score, "
    "severity score, temporal information (JSONB), machine learning risk label and confidence, "
    "disruption probability, impact score, sentiment label, and sentiment score."
)

body(
    "The ForecastSnapshot table is a dedicated table that stores the output of the Two-Stage "
    "XGBoost forecasting engine. Each record is associated with a supplier node, a forecast "
    "origin date, and a horizon day offset. It stores the predicted risk value (yhat), the "
    "lower and upper prediction intervals (yhat_lower, yhat_upper), the actual realised risk "
    "value if available (y_actual), and the forecast method identifier. This table enables "
    "the system to serve historical and on-demand 14-day risk forecasts through the API."
)

body(
    "The ResilienceHistory table stores the historical risk scores of every supplier over time. "
    "Each record is identified by a unique ID and a foreign key connecting it to the Supplier "
    "table. The attributes include the recorded date and time and the risk score. The system "
    "uses this table to display historical performance trends of suppliers on the interface."
)

# 3.3.4
heading2("3.3.4\tUse Case Diagram")

placeholder_figure("Figure 3.7 – Insert Use Case Diagram")
fig_caption("Figure 3.7 Use Case Diagram")

use_cases = [
    ("1. View Map Dashboard",
     ["The user can access the main dashboard of the system.",
      "The system displays a world map with supplier nodes.",
      "Each node represents a supplier and is colour-coded based on its current risk level."]),
    ("2. Filter Suppliers by Supply Chain or Product",
     ["The user can filter the displayed supplier nodes by selecting a specific supply chain or product.",
      "The system will update the world map to show only the relevant suppliers."]),
    ("3. View Supplier Details (Risk and Resilience)",
     ["The user can click on any supplier node on the map.",
      "The system will display detailed information including current risk score, risk level, criticality, and latest related disruption events."]),
    ("4. View Resilience History",
     ["The user can view the historical risk trend of a selected supplier.",
      "The system will display a time-based chart showing changes in risk scores.",
      "This allows the user to analyse long-term supplier performance."]),
    ("5. View 14-Day Risk Forecast",
     ["The user can view the 14-day forward risk forecast for a selected supplier.",
      "The system will display a forecast chart generated by the Two-Stage XGBoost model, showing predicted risk values with upper and lower confidence intervals.",
      "This allows the user to anticipate upcoming disruption risks before they occur."]),
    ("6. Search Supplier",
     ["The user can search for a supplier by name or location.",
      "The system will highlight or display the matching supplier on the map and in the detail panel."]),
    ("7. View Related News Articles",
     ["The user can view the news articles associated with a supplier or disruption event.",
      "The system displays article titles, sources, and publication dates."]),
    ("8. Trigger Manual Refresh",
     ["The administrator can manually trigger the system to fetch the latest news articles from the RSS feeds.",
      "The system will immediately run the inference pipeline and update the database."]),
    ("9. Manage Supplier Records",
     ["The administrator can add new suppliers to the system.",
      "The administrator can update or remove existing supplier information."]),
    ("10. Manage Supply Chain Records",
     ["The administrator can add, update, or remove supply chain or product records."]),
    ("11. View System Logs and Reports",
     ["The administrator can view system activity logs and processing reports."]),
]

for title, bullets in use_cases:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    for b in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        run = bp.add_run(b)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        bp.paragraph_format.left_indent  = Inches(0.5)
        bp.paragraph_format.space_after  = Pt(2)

doc.add_paragraph()

# 3.3.5
heading2("3.3.5\tUse Case Description")

uc_tables = [
    ("Table 3.3 View Map Dashboard Use Case Description",
     [("Use Case Name","View Map Dashboard"),
      ("Description","User views the main dashboard showing supplier nodes on world map."),
      ("Pre-Condition(s)","System is running."),
      ("Post-Condition(s)","Dashboard is displayed."),
      ("Actor","User"),
      ("Primary Flow","- Open system homepage\n- System loads map and supplier nodes"),
      ("Alternative / Exception Flow","If data fails to load, error message is shown.")]),
    ("Table 3.4 Filter Suppliers by Supply Chain Use Case Description",
     [("Use Case Name","Filter Suppliers by Supply Chain"),
      ("Description","User filters suppliers by product or supply chain."),
      ("Pre-Condition(s)","Dashboard is open."),
      ("Post-Condition(s)","Filtered nodes are shown."),
      ("Actor","User"),
      ("Primary Flow","- Select filter\n- System queries database\n- Map updates"),
      ("Alternative / Exception Flow","If no results, empty map is shown.")]),
    ("Table 3.5 View Supplier Details Use Case Description",
     [("Use Case Name","View Supplier Details"),
      ("Description","User views detailed supplier risk and criticality information."),
      ("Pre-Condition(s)","Map is visible."),
      ("Post-Condition(s)","Detail panel is shown."),
      ("Actor","User"),
      ("Primary Flow","- Click supplier node\n- System fetches data\n- Display details including risk score, criticality, and related events"),
      ("Alternative / Exception Flow","If data missing, error shown.")]),
    ("Table 3.6 View Resilience History Use Case Description",
     [("Use Case Name","View Resilience History"),
      ("Description","User views historical risk trend for a supplier."),
      ("Pre-Condition(s)","Supplier page open."),
      ("Post-Condition(s)","History chart shown."),
      ("Actor","User"),
      ("Primary Flow","- Click history tab\n- System loads historical risk records\n- Display chart"),
      ("Alternative / Exception Flow","If no data, empty chart shown.")]),
    ("Table 3.7 View 14-Day Risk Forecast Use Case Description",
     [("Use Case Name","View 14-Day Risk Forecast"),
      ("Description","User views the Two-Stage XGBoost 14-day forward risk forecast for a selected supplier."),
      ("Pre-Condition(s)","Supplier page open."),
      ("Post-Condition(s)","Forecast chart displayed with predicted risk values and confidence intervals."),
      ("Actor","User"),
      ("Primary Flow","- Click forecast tab\n- System retrieves or generates forecast snapshot\n- Display forecast chart with yhat, yhat_lower, yhat_upper for 14 days"),
      ("Alternative / Exception Flow","If forecast snapshot is unavailable, system generates it on demand.")]),
    ("Table 3.8 Search Supplier Use Case Description",
     [("Use Case Name","Search Supplier"),
      ("Description","User searches supplier by name."),
      ("Pre-Condition(s)","Dashboard open."),
      ("Post-Condition(s)","Matching suppliers shown."),
      ("Actor","User"),
      ("Primary Flow","- Enter keyword\n- System searches database\n- Show result"),
      ("Alternative / Exception Flow","If not found, show no result.")]),
    ("Table 3.9 View Related News Articles Use Case Description",
     [("Use Case Name","View Related News Articles"),
      ("Description","User views news related to supplier."),
      ("Pre-Condition(s)","Supplier page open."),
      ("Post-Condition(s)","News list displayed."),
      ("Actor","User"),
      ("Primary Flow","- Click news tab\n- System queries database\n- Display articles"),
      ("Alternative / Exception Flow","If none, empty list shown.")]),
    ("Table 3.10 Trigger Manual Refresh Use Case Description",
     [("Use Case Name","Trigger Manual Refresh"),
      ("Description","Admin manually triggers RSS news update."),
      ("Pre-Condition(s)","Admin logged in."),
      ("Post-Condition(s)","System fetches from RSS feeds and updates data."),
      ("Actor","Admin"),
      ("Primary Flow","- Click refresh\n- System polls RSS feeds\n- Run inference pipeline\n- Update database and forecast snapshots"),
      ("Alternative / Exception Flow","If feed unreachable, show error.")]),
    ("Table 3.11 Manage Supplier Nodes Use Case Description",
     [("Use Case Name","Manage Supplier Nodes"),
      ("Description","Admin manages suppliers."),
      ("Pre-Condition(s)","Admin logged in."),
      ("Post-Condition(s)","Supplier database updated."),
      ("Actor","Admin"),
      ("Primary Flow","- Add/Edit/Delete supplier\n- System validates\n- Save to database"),
      ("Alternative / Exception Flow","If invalid data, reject input.")]),
    ("Table 3.12 Manage Supply Chains Use Case Description",
     [("Use Case Name","Manage Supply Chains"),
      ("Description","Admin manages supply chains."),
      ("Pre-Condition(s)","Admin logged in."),
      ("Post-Condition(s)","Supply chain database updated."),
      ("Actor","Admin"),
      ("Primary Flow","- Add/Edit/Delete supply chain\n- Save changes"),
      ("Alternative / Exception Flow","If conflict, show error.")]),
    ("Table 3.13 View System Logs / Reports Use Case Description",
     [("Use Case Name","View System Logs / Reports"),
      ("Description","Admin views system logs and known issues."),
      ("Pre-Condition(s)","Admin logged in."),
      ("Post-Condition(s)","Logs displayed."),
      ("Actor","Admin"),
      ("Primary Flow","- Open logs page\n- System loads logs"),
      ("Alternative / Exception Flow","If no logs, show empty page.")]),
]

for caption, rows in uc_tables:
    table_caption(caption)
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, (label, val) in enumerate(rows):
        cells = t.rows[r_idx].cells
        run = cells[0].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        run2 = cells[1].paragraphs[0].add_run(val)
        run2.font.size = Pt(11)
        run2.font.name = "Times New Roman"
    doc.add_paragraph()

# 3.3.6
heading2("3.3.6\tUser Interface")

body(
    "This section describes the user interface design of the Supply Chain Resilience "
    "Monitoring System. The interface was designed to present supply chain risk, resilience, "
    "disruption information, and forward-looking forecasts in a clear and structured manner "
    "to support effective monitoring. The following figures illustrate the mock user interface "
    "and explain the functionality of each interface."
)

placeholder_figure("Figure 3.8 – Insert Map Dashboard Interface screenshot")
fig_caption("Figure 3.8 Map Dashboard Interface")

body(
    "Figure 3.8 displays the Map Dashboard, which gives a global view of the supply chain. "
    "Supplier nodes are presented on an interactive world map, and their current level of risk "
    "is colour-coded, providing the user with an opportunity to easily locate high-risk suppliers. "
    "The most important metrics such as the total number of suppliers, high-risk suppliers, "
    "average risk score, and active alerts are presented in summary indicators at the top of "
    "the screen. Suppliers can also be filtered by supply chain or product using filters which "
    "allow users to narrow down the view."
)

body(
    "Selecting a supplier node reveals a detail panel where the supplier's location, risk score, "
    "criticality level, the supply chains that include this supplier, and recent disruption events "
    "can be seen. This interface facilitates fast monitoring and top-level decision making."
)

placeholder_figure("Figure 3.9 – Insert Supplier Details Interface screenshot")
fig_caption("Figure 3.9 Supplier Details Interface")

body(
    "Figure 3.9 shows the Suppliers screen, in which all monitored suppliers are displayed in "
    "a table format. The table shows supplier location, supply chains, risk score, criticality, "
    "and status. It also has a search feature and a filtering system so that the user can find "
    "particular suppliers faster. The visual representation of risk values through progress bars "
    "and colour-coded labels allows suppliers to be easily compared."
)

placeholder_figure("Figure 3.10 – Insert Resilience History Interface screenshot")
fig_caption("Figure 3.10 Resilience History Interface")

body(
    "Figure 3.10 demonstrates the Resilience History screen that enables users to track the "
    "past performance of an identified supplier. The interface shows the current risk score and "
    "overall trend status. Changes in risk scores across time are represented in line charts. "
    "This screen is useful for long-term analysis of stability and recovery trends of suppliers."
)

placeholder_figure("Figure 3.11 – Insert 14-Day Forecast Interface screenshot")
fig_caption("Figure 3.11 14-Day Risk Forecast Interface")

body(
    "Figure 3.11 shows the 14-Day Risk Forecast screen, which visualises the output of the "
    "Two-Stage XGBoost forecasting engine for a selected supplier. The interface displays a "
    "forecast chart showing the predicted daily risk values (yhat) and the associated "
    "confidence intervals (yhat_lower and yhat_upper) over the next 14 days. Where historical "
    "data is available, the actual realised risk values are overlaid on the chart to allow "
    "comparison against forecast accuracy. This screen enables operators to anticipate "
    "upcoming disruption risks before they materialise."
)

placeholder_figure("Figure 3.12 – Insert News & Events Interface screenshot")
fig_caption("Figure 3.12 News & Events Interface")

body(
    "Figure 3.12 provides an overview of the News and Events screen, which displays news "
    "articles and disruption events of suppliers under observation. The associated news section "
    "includes article sources, date of publication, headlines, and machine learning risk labels. "
    "The disruption events section displays organised events with corresponding risk levels, "
    "sentiment scores, and status. This interface provides users with the ability to comprehend "
    "the real-world events driving supplier risk scores, which facilitates transparency and "
    "explainable system output."
)

placeholder_figure("Figure 3.13 – Insert Administration Interface screenshot")
fig_caption("Figure 3.13 Administration Interface")

body(
    "Figure 3.13 shows the Administration screen, intended to monitor and control the system. "
    "The administrators may manually initiate RSS news updates, check database status, system "
    "health, and system logs. The system log table documents important operations including "
    "RSS ingestion, risk calculation, forecast generation, and data synchronisation, along with "
    "their execution state. This screen aids in system maintenance and operational control."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.4  Phase 3: Development
# ══════════════════════════════════════════════════════════════════════════════
heading1("3.4\tPhase 3: Development")

body(
    "Phase 3 entailed the implementation of the supply chain disruption detection system as "
    "a local prototype web application. This stage combined the frontend, backend, and machine "
    "learning modules into a working solution. The tasks involved in development were to create "
    "a ReactJS front-end interface, a Python back-end API with the FastAPI framework, a "
    "multi-model ML inference engine combining XGBoost, FinBERT, and GLiNER2, RSS-based "
    "live data ingestion, and a Two-Stage XGBoost forecasting engine for 14-day risk "
    "prediction. The aim was to make sure that real-time news articles that may act as early "
    "warning signs of supply chain disruptions are automatically downloaded, processed "
    "through the ML pipeline, and displayed on the web interface with relevant risk, "
    "sentiment, and forecast information."
)

# 3.4.1
heading2("3.4.1\tHardware and Software Requirements")

body(
    "The hardware requirements outline the required minimum of computing resources needed "
    "to support system development, model training, and real-time inference processes. "
    "Large news data volumes and machine learning workloads require sufficient processing "
    "capacity, memory, and storage. The system has been initially set up to run in a local "
    "machine environment in order to enable controlled development and testing."
)

add_table(
    ["Component", "Specification", "Justification"],
    [
        ["Processor (CPU)", "64-bit multi-core processor (e.g., Intel Core i5 or AMD equivalent, 2.5 GHz or above).",
         "A multi-core CPU helps in handling the web server and ML inference in parallel."],
        ["Memory (RAM)", "8 GB RAM minimum (16 GB recommended).",
         "Sufficient memory is needed to load FinBERT and GLiNER2 models and handle data processing for multiple news articles simultaneously without performance degradation."],
        ["Storage", "At least 5–10 GB free disk space.",
         "This accommodates the project source code, installed libraries, cached ML models (FinBERT ~400 MB, GLiNER2 ~200 MB), PostgreSQL data, and pipeline outputs."],
        ["Graphics (GPU)", "NVIDIA GPU with CUDA support (e.g., NVIDIA GTX/RTX series) or Apple Silicon GPU (MPS) for accelerating model inference.",
         "The system is capable of running on CPU alone, but a GPU can significantly speed up FinBERT and GLiNER2 inference when processing large batches of articles. The system auto-detects the available device."],
        ["Network", "Active internet connection for the backend to poll RSS feeds and for accessing external geocoding services.",
         "Also used during development for installing packages (via npm/pip) and for online dataset access."],
        ["Operating System", "Windows 10 (64-bit), Linux Ubuntu 20.04+, or macOS.",
         "The development stack (Node, Python, PostgreSQL) is cross-platform. macOS users require OpenMP (brew install libomp) for XGBoost."],
    ],
    "Table 3.14 Hardware Requirements"
)

body(
    "Software requirements indicate the programming environment, libraries, and tools required "
    "to develop and implement the proposed supply chain resilience analysis system. These "
    "software modules facilitate data processing, machine learning model development, "
    "integration of live news data, and interactive visualisation of supply chain networks."
)

add_table(
    ["Software", "Version", "Purpose"],
    [
        ["Python", "3.11 or higher", "Core programming environment for backend development, data processing, and machine learning model execution."],
        ["FastAPI", "Latest stable version", "RESTful backend API framework; handles communication between frontend and ML components."],
        ["ReactJS (Vite)", "Latest stable version", "Builds the interactive web-based frontend dashboard for visualising supply chain disruption, resilience, and forecast data."],
        ["PostgreSQL", "Latest stable version", "Relational database that serves as the system of record for supplier data, events, risk history, and forecast snapshots."],
        ["XGBoost", "Latest version", "Implements the gradient boosting models for tri-class headline risk classification, disruption detection, impact regression, and the Two-Stage risk forecaster."],
        ["Hugging Face Transformers (FinBERT)", "Latest version", "Provides FinBERT (ProsusAI/finbert) for financial domain sentiment analysis on news articles, producing sentiment_label and sentiment_score."],
        ["GLiNER2", "fastino/gliner2-base-v1", "Named entity recognition model for extracting geographic locations from news text. Replaced the earlier spaCy-based NER approach."],
        ["spaCy", "Latest version (en_core_web_sm)", "Used exclusively for temporal date extraction to identify when an event is expected to occur. Not used for NER in the current system."],
        ["Pandas & NumPy", "Latest version", "Handle data preprocessing, numerical computation, and feature engineering throughout the pipeline."],
        ["Node.js & npm", "Latest LTS version", "Supports ReactJS frontend development and dependency management."],
    ],
    "Table 3.15 Software Requirements"
)

# 3.4.2
heading2("3.4.2\tMachine Learning Model Pseudocode")

body(
    "The subsection below details the pseudocode of the machine learning components applied "
    "in the Supply Chain Resilience Monitoring System. The pseudocode outlines the offline "
    "training process and the online inference process, which together constitute the multi-model "
    "machine learning architecture. This pseudocode is meant to show in detail the logical "
    "sequence of events in disruption detection, model deployment, risk score update, and "
    "14-day forecast generation."
)

heading3("3.4.2.1\tAlgorithm 1: Offline Model Training Procedure")

pseudocode_block([
    "Input:  Historical news dataset from Hugging Face",
    "Output: Trained disruption classification model (classifier.pkl),",
    "        disruption classifier (disruption_classifier.pkl),",
    "        impact regressor (impact_regressor_v2.pkl),",
    "        Two-Stage XGBoost forecaster (forecast_event_prob.json,",
    "                                      forecast_severity_q75.json)",
    "",
    "1.  Begin",
    "2.  Load historical news dataset",
    "3.  Extract text and label fields from dataset",
    "",
    "4.  For each news article in dataset do",
    "5.      Clean text (remove punctuation, numbers, special characters)",
    "6.      Convert text to lowercase",
    "7.      Perform tokenization",
    "8.      Remove stop words",
    "9.      Apply lemmatization",
    "10. End For",
    "",
    "11. Apply TF-IDF vectorisation to produce feature vectors",
    "12. Encode tri-class labels (LOW, MEDIUM, HIGH) using LabelEncoder",
    "13. Split dataset into training and validation sets",
    "",
    "14. Train XGBoost headline classifier (tri-class: LOW/MEDIUM/HIGH)",
    "15. Train XGBoost binary disruption classifier",
    "16. Train XGBoost impact regressor (predicted_impact_score)",
    "17. Train Two-Stage XGBoost forecaster:",
    "        Stage 1 — XGBClassifier to predict P(event) per day",
    "        Stage 2 — XGBRegressor (quantile q=0.75) to predict E[severity | event]",
    "        Freeze-window: all features computed from actual data at forecast_date;",
    "                       day_offset encodes the horizon position",
    "",
    "18. Evaluate each model using validation data",
    "19. Compute accuracy, precision, recall, F1-score (classifiers)",
    "       Compute MAE and RMSE (regressors and forecaster)",
    "20. If model performance is satisfactory then",
    "21.     Save trained models to disk",
    "22.     Store model metadata (version, date, metrics)",
    "23. End If",
    "24. End",
])

body(
    "Algorithm 1 describes how the disruption detection and forecasting models are trained "
    "using historical news data. It begins with loading the news dataset and extracting the "
    "text and labels. News articles are cleaned and preprocessed through tokenization, stop "
    "word removal, and lemmatization. TF-IDF vectorisation then converts the cleaned text into "
    "numerical features for the XGBoost classifiers. Three separate XGBoost models are "
    "trained: a tri-class headline risk classifier, a binary disruption classifier, and an "
    "impact regressor."
)

body(
    "The Two-Stage XGBoost forecaster is trained separately using historical risk snapshots. "
    "Stage 1 predicts the probability that a disruption event occurs on a given future day, "
    "while Stage 2 predicts the expected severity given that an event occurs. A freeze-window "
    "architecture is used so that all features are computed from actual data at the forecast "
    "origin date, with a day_offset feature distinguishing horizon positions. This prevents "
    "data leakage and ensures the model can be evaluated on truly out-of-sample future days. "
    "All models and their metadata are stored to disk upon reaching satisfactory performance."
)

heading3("3.4.2.2\tAlgorithm 2: Online News Inference and Risk Update Procedure")

pseudocode_block([
    "Input:  Live news articles from RSS feeds",
    "Output: Updated supplier risk scores, disruption event records,",
    "        refreshed 14-day forecast snapshots",
    "",
    "1.  Begin",
    "2.  Load trained models (classifier.pkl, disruption_classifier.pkl,",
    "                          impact_regressor_v2.pkl, forecast_event_prob.json,",
    "                          forecast_severity_q75.json)",
    "3.  Poll configured RSS feeds and retrieve latest articles",
    "",
    "4.  For each retrieved news article do",
    "5.      Validate article and remove duplicates",
    "6.      Clean and preprocess article text",
    "",
    "7.      Extract named location entities using GLiNER2",
    "8.      Geocode extracted locations",
    "9.      Match geocoded locations with supplier node database",
    "",
    "10.     Classify article headline using XGBoost tri-class classifier",
    "        → ml_risk_label (LOW / MEDIUM / HIGH)",
    "        → ml_risk_probabilities",
    "        → risk_score (mapped from label)",
    "",
    "11.     Compute sentiment using FinBERT",
    "        → sentiment_label, sentiment_score",
    "",
    "12.     Predict disruption probability using XGBoost disruption classifier",
    "        → predicted_disruption_probability",
    "",
    "13.     Predict impact score using XGBoost impact regressor",
    "        → predicted_impact_score",
    "",
    "14.     If related supplier node is found then",
    "15.         Store event record in database with all ML outputs",
    "",
    "16.         If disruption probability exceeds threshold then",
    "17.             Create disruption event record",
    "18.             Recompute supplier exposure index (0–100):",
    "                strength = min(100, predicted_impact_score / 3)",
    "                exposure = min(100, 0.62 * avg(strength) + 0.38 * max(strength))",
    "19.             Append record to resilience history",
    "20.             Regenerate 14-day forecast snapshot via Two-Stage XGBoost:",
    "                yhat = P(event) * severity  for each of 14 horizon days",
    "                Store result in forecast_snapshots table",
    "21.         Else",
    "22.             Store article as non-disruption record",
    "23.         End If",
    "24.     End If",
    "25. End For",
    "",
    "26. Update frontend API with latest risk, resilience, and forecast data",
    "27. End",
])

body(
    "Algorithm 2 describes the use of the trained models during system operation. The RSS "
    "ingestion module polls configured RSS feeds and retrieves live news articles, which are "
    "validated and deduplicated before processing. Named entity recognition for locations is "
    "performed using the GLiNER2 model (fastino/gliner2-base-v1), which replaces the earlier "
    "spaCy-based NER approach to provide more accurate geographic entity extraction."
)

body(
    "Once a supplier association is identified, the article is passed through the multi-model "
    "inference pipeline. The XGBoost headline classifier assigns a tri-class risk label. FinBERT "
    "computes a financial domain sentiment score. The disruption classifier and impact regressor "
    "produce a disruption probability and an impact score respectively. When a disruption is "
    "detected, the system generates a disruption event record, recomputes the supplier exposure "
    "index using a weighted blend of average and maximum event strength, and triggers the "
    "Two-Stage XGBoost forecasting engine to regenerate the 14-day risk forecast for the "
    "affected node. Non-disruption articles are stored without impacting supplier scores."
)

# 3.4.3
heading2("3.4.3\tRSS Feed Ingestion")

body(
    "RSS feed ingestion is incorporated into the system as the primary external source of "
    "real-time news articles concerning supply chain disruptions. The RSS feeds are configured "
    "in a JSON file (config/rss_feeds.json) as an array of feed entries, each containing a "
    "URL and a source label. The system supports any number of feed sources, providing "
    "broad coverage of international supply chain news."
)

body(
    "The ingestion module (src/rss_ingest.py) is implemented in Python and is triggered on a "
    "scheduled polling interval (default approximately 10 minutes) by a background worker "
    "managed by the FastAPI lifespan handler. An administrator may also manually trigger a "
    "single ingestion cycle through the POST /admin/rss-ingest/trigger API endpoint."
)

body(
    "Upon each polling cycle, the system parses all configured RSS feeds, extracts article "
    "URLs, titles, summaries, and publication timestamps, and validates each article against "
    "previously ingested records to avoid duplicate processing. Articles that pass validation "
    "are immediately passed to the preprocessing and ML inference pipeline, providing "
    "near-real-time disruption detection. Following inference, the results are persisted to the "
    "database and the forecast snapshots are refreshed for any affected supplier nodes."
)

placeholder_figure("Figure 3.14 – Insert RSS ingestion code segment")
fig_caption("Figure 3.14 Coding Segment for RSS Feed Ingestion")

# 3.4.4  (was 3.4.4 Limitations of NewsAPI — now GLiNER2 NER)
heading2("3.4.4\tGLiNER2 Location Named Entity Recognition")

body(
    "Location named entity recognition (NER) is a critical step in associating news articles "
    "with the correct supplier nodes. The system uses GLiNER2 (fastino/gliner2-base-v1), a "
    "lightweight generalised NER model based on the GLiNER architecture, for this purpose. "
    "This model was introduced to replace an earlier BERT-Large based NER approach, "
    "significantly reducing inference memory requirements and batch processing time while "
    "maintaining high extraction accuracy for geographic location entities."
)

body(
    "GLiNER2 operates by performing span-level entity classification on the input text given "
    "a set of target entity type labels (in this case, location). The model is invoked in batch "
    "mode (extract_locations_batch in src/preprocessing.py) to process multiple articles "
    "efficiently. The GLINER_MODEL environment variable allows the model identifier to be "
    "changed without modifying the source code, supporting future model upgrades."
)

body(
    "The extracted location entities are subsequently passed to a geocoding module that "
    "resolves them to latitude and longitude coordinates, which are then matched against the "
    "canonical supplier node list to identify the relevant supplier for each article."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.4.5  Two-Stage XGBoost Risk Forecasting
# ══════════════════════════════════════════════════════════════════════════════
heading2("3.4.5\tTwo-Stage XGBoost Risk Forecasting")

body(
    "The system incorporates a dedicated risk forecasting component that generates 14-day "
    "forward risk predictions for each supplier node. The forecasting engine is implemented "
    "as a Two-Stage XGBoost model (src/forecast_snapshots.py), which forms the primary "
    "production forecasting path for all forecast API endpoints and pipeline runs."
)

body(
    "The design of the forecaster is motivated by the need to decompose the risk prediction "
    "task into two interpretable sub-problems. Stage 1 estimates the probability that a "
    "disruption event occurs on a given future day (P(event)), trained as an XGBoost binary "
    "classifier and stored in models/forecast_event_prob.json. Stage 2 estimates the expected "
    "severity of the disruption given that an event does occur (E[severity | event]), trained "
    "as an XGBoost quantile regressor at the 75th percentile (q=0.75) and stored in "
    "models/forecast_severity_q75.json. A mean regression variant "
    "(models/forecast_severity_mean.json) is also available for comparison. The final "
    "predicted risk value for each horizon day is computed as:"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("yhat  =  P(event)  ×  E[severity | event]")
run.bold = True
run.font.size = Pt(12)
run.font.name = "Times New Roman"
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)

body(
    "A key architectural decision is the use of a freeze-window approach. Rather than "
    "generating features recursively (i.e., using predictions from earlier horizon days as "
    "inputs to later days, which introduces compounding error), all features for every horizon "
    "day are computed entirely from actual data available at the forecast origin date "
    "(forecast_date). A day_offset feature (integer from 1 to 14) is included as an explicit "
    "model input to differentiate each horizon position. This design eliminates recursive "
    "data leakage and allows the model to be evaluated on truly out-of-sample future days."
)

body(
    "The forecaster computes approximately 31 features per horizon day, encompassing "
    "frozen features (derived from actual data at forecast_date, constant across the horizon) "
    "and day-varying features (the day_offset and any day-specific calendar features). "
    "Training is performed via scripts/train_two_stage_forecast.py and the trained models "
    "are stored as JSON artifacts in the models/ directory."
)

body(
    "Forecast results are persisted in the forecast_snapshots table of the PostgreSQL "
    "database. Each snapshot record stores the node name, forecast origin date, horizon "
    "day offset, predicted risk value (yhat), lower and upper prediction intervals "
    "(yhat_lower, yhat_upper), and the actual realised risk value where available (y_actual). "
    "The forecast API endpoint (GET /suppliers/{node_name}/forecast) returns these "
    "snapshots, generating them on demand if a pre-computed snapshot for the requested "
    "date is not found."
)

body(
    "The method key exposed through the API is 'xgboost' for the default q75 production "
    "model and 'xgboost_mean' for the mean regression comparison variant. Legacy EDSF/Prophet "
    "forecasting code is retained in src/predictive_forecasting.py but is no longer on any "
    "active production path."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.5  Phase 4: Testing
# ══════════════════════════════════════════════════════════════════════════════
heading1("3.5\tPhase 4: Testing")

body(
    "This stage aims at ensuring that the proposed Supply Chain Resilience Monitoring System "
    "is functioning correctly, is reliable, and is working well. The test is done to verify "
    "that every component of the system is functioning as per the requirements outlined, "
    "such as data ingestion, machine learning inference, risk and resilience score updates, "
    "forecast snapshot generation, and user interface interactions. The test scenarios are "
    "modelled on functional requirements and implemented in a local deployment environment."
)

heading2("3.5.1\tFunctional Testing")

body(
    "Table 3.16 presents the functional test cases conducted for the system. Each test case "
    "includes the objective, preconditions, testing steps, expected results, actual results, "
    "and test status. All tests were executed after successful system setup and model loading."
)

add_table(
    ["Test Case ID", "Objective", "Precondition", "Steps", "Expected Result", "Actual Result", "Status"],
    [
        ["TC001", "Verify that machine learning models are loaded successfully",
         "Application server is launched", "Start the backend API server",
         "All trained ML models (classifier.pkl, disruption_classifier.pkl, impact_regressor_v2.pkl, GLiNER2, FinBERT, XGBoost forecaster) are loaded without runtime errors",
         "Models loaded successfully without errors", "Pass"],
        ["TC002", "Verify real-time RSS news ingestion and disruption classification",
         "System is running and RSS feeds are configured in config/rss_feeds.json",
         "Trigger RSS ingestion process from API",
         "Relevant news articles are retrieved from RSS feeds and classified correctly with tri-class risk labels",
         "News articles processed and classified successfully", "Pass"],
        ["TC003", "Verify supplier risk score update after disruption detection",
         "Disruption event exists in the database",
         "Execute risk scoring module",
         "Supplier exposure index (0–100) is updated based on the detected event using the weighted blend formula",
         "Risk score updated correctly in database", "Pass"],
        ["TC004", "Verify resilience score computation and historical update",
         "Risk score has been updated",
         "Run resilience score calculation",
         "Risk exposure index is recalculated and stored in resilience history table",
         "Resilience history updated successfully", "Pass"],
        ["TC005", "Verify 14-day risk forecast generation",
         "Supplier node exists and ML models are loaded",
         "Call GET /suppliers/{node_name}/forecast",
         "14-day forecast snapshot is returned with yhat, yhat_lower, yhat_upper for each horizon day",
         "Forecast generated successfully and returned by API", "Pass"],
        ["TC006", "Verify world map visualisation of supplier nodes",
         "Frontend application is running",
         "Open world map homepage",
         "Supplier nodes and links are displayed correctly on map with colour-coded risk levels",
         "Map rendered with correct supplier visualisation", "Pass"],
        ["TC007", "Verify supplier filtering by supply chain",
         "Supplier and product data exist",
         "Select a specific product supply chain filter",
         "Only suppliers related to the selected supply chain are shown",
         "Correct suppliers filtered and displayed", "Pass"],
        ["TC008", "Verify supplier detail page display",
         "Supplier record exists",
         "Click on a supplier node",
         "Supplier details including risk score, criticality, and historical data are displayed",
         "Supplier detail page displayed correctly", "Pass"],
    ],
    "Table 3.16 Functional Testing"
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.6  Summary
# ══════════════════════════════════════════════════════════════════════════════
heading1("3.6\tSummary")

body(
    "This chapter has presented the complete methodology adopted for the development of "
    "the Supply Chain Resilience Monitoring System. The overall development process was "
    "guided by the Waterfall model, which was selected due to its structured, sequential "
    "nature and suitability for projects with well-defined requirements. The methodology "
    "clearly outlined each development phase, beginning from system requirement analysis "
    "and design, followed by implementation, testing, and deployment. This method "
    "guaranteed systematic progress and explicit validation at every stage of development."
)

body(
    "The system architecture was modelled as a hybrid system that integrates offline machine "
    "learning model training with online inference. The offline training phase produced "
    "multiple complementary models: an XGBoost tri-class headline risk classifier, a FinBERT "
    "sentiment analyser, an XGBoost disruption classifier, an XGBoost impact regressor, and "
    "a Two-Stage XGBoost risk forecaster employing a freeze-window architecture to generate "
    "14-day forward risk predictions. Online inference processes live news articles retrieved "
    "through a scheduled RSS ingestion pipeline, with location entity extraction performed "
    "by the GLiNER2 model."
)

body(
    "The development stage aimed at deploying a backend service based on Python and FastAPI, "
    "integrating all trained machine learning models into a unified inference pipeline, "
    "persisting all outputs into a PostgreSQL database, and creating a web interface based on "
    "ReactJS (Vite) to visualise the data. The system offers interactive capabilities including "
    "a world map with colour-coded supplier node view, supply chain relationship view, risk "
    "and resilience score display, historical trend view, and a 14-day forward risk forecast "
    "chart per supplier. All these elements combine to provide timely information on supply "
    "chain vulnerability and resilience."
)

body(
    "Finally, this chapter has shown a methodological approach towards developing the proposed "
    "system. By integrating multi-model machine learning, named entity recognition, financial "
    "sentiment analysis, forward-looking risk forecasting, and interactive visualisation within "
    "a well-defined methodological framework, the Supply Chain Resilience Monitoring System "
    "provides a practical tool for enhancing supply chain risk awareness and resilience "
    "assessment. The methodology adopted in this chapter establishes a strong foundation for "
    "the implementation results and evaluation discussed in the subsequent chapters."
)

# ── save ─────────────────────────────────────────────────────────────────────
out = "/Users/meordanish/Desktop/Projects/SupplyChainForecast/docs/chap3_updated.docx"
doc.save(out)
print(f"Saved: {out}")
